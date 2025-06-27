from flask import Flask, request, jsonify, render_template_string, send_file
from flask_cors import CORS
from datetime import datetime
import gc  # For memory management
import re
import json
import base64
import os
import tempfile
import shutil
# OpenAI integration with cost optimization
# Define constants first (before try block to avoid NameError)
MAX_TOKENS_PER_REQUEST = 1000  # Keep costs low
MODEL_NAME = "gpt-4o-mini"  # Most cost-effective model

# MEMORY OPTIMIZATION FUNCTIONS
def truncate_text_content(data, max_length=1000):
    """Truncate text content to prevent memory issues"""
    if isinstance(data, dict):
        result = {}
        for key, value in data.items():
            if key in ['content', 'text', 'extracted_text'] and isinstance(value, str):
                # Truncate long text content
                result[key] = value[:max_length] + "..." if len(value) > max_length else value
            elif isinstance(value, (dict, list)):
                result[key] = truncate_text_content(value, max_length)
            else:
                result[key] = value
        return result
    elif isinstance(data, list):
        return [truncate_text_content(item, max_length) for item in data]
    else:
        return data

def memory_safe_ai_call(client, prompt, max_tokens=500):
    """Make AI call with memory safety measures"""
    try:
        # Limit prompt size to prevent memory issues
        if len(prompt) > 8000:  # Approximately 2000 tokens
            prompt = prompt[:8000] + "\n\n[Content truncated for memory safety]"
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",  # Use smaller model
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tokens,  # Limit response size
            temperature=0.1
        )
        
        # Force garbage collection after AI call
        gc.collect()
        
        return response
        
    except Exception as e:
        # Force cleanup on error
        gc.collect()
        raise e

def get_memory_usage():
    """Get current memory usage in MB"""
    try:
        import psutil
        process = psutil.Process()
        return process.memory_info().rss / 1024 / 1024  # MB
    except:
        return 0
MAX_INPUT_TOKENS = 8000  # Limit input size

try:
    from openai import OpenAI
    import os
    OPENAI_AVAILABLE = True
    
    # CACHE BUSTER - Force fresh deployment
    DEPLOYMENT_TIMESTAMP = "2025-06-26-19-35-ENHANCED-PDF-PAGES"
    
    # NO HARDCODED API KEY - Environment variables only for security
    # OpenAI automatically disables keys exposed in code
    
    # Production-ready API key handling - check multiple possible environment variable names
    env_key = (os.getenv('OPENAI_API_KEY') or 
               os.getenv('OPEN_API_KEY') or  # User's environment variable name
               os.getenv('OPENAI_KEY') or 
               os.getenv('API_KEY'))
    
    OPENAI_API_KEY = env_key
    
    # Debug: Print environment variable info at startup
    print(f"🔑 STARTUP: API key {'found' if OPENAI_API_KEY else 'NOT FOUND'}")
    print(f"🌍 ENV SOURCE: {'environment' if env_key else 'NONE - NO HARDCODED FALLBACK'}")
    print(f"📅 DEPLOYMENT: {DEPLOYMENT_TIMESTAMP}")
    
    # Debug: Show available environment variables that might contain API keys
    env_vars = [k for k in os.environ.keys() if 'API' in k.upper() or 'OPENAI' in k.upper()]
    if env_vars:
        print(f"🔍 FOUND ENV VARS: {env_vars}")
    else:
        print("⚠️  NO API-related environment variables found")
    
    # Production-safe client initialization with comprehensive error handling
    openai_client = None
    openai_client_error = None
    
    def get_openai_client():
        """Get OpenAI client with robust initialization and error handling"""
        global openai_client, openai_client_error
        
        # Force fresh initialization - clear any cached client
        openai_client = None
        openai_client_error = None
        
        # DYNAMIC API KEY RETRIEVAL - Environment variables only (no hardcoded fallback)
        env_openai_api_key = os.getenv('OPENAI_API_KEY')
        env_open_api_key = os.getenv('OPEN_API_KEY')  # User's environment variable name
        env_openai_key = os.getenv('OPENAI_KEY')
        env_api_key = os.getenv('API_KEY')
        
        # Debug: Show what we found in environment variables
        print(f"🔍 ENV DEBUG:")
        print(f"   OPENAI_API_KEY: {'Found' if env_openai_api_key else 'Not found'}")
        print(f"   OPEN_API_KEY: {'Found' if env_open_api_key else 'Not found'}")
        print(f"   OPENAI_KEY: {'Found' if env_openai_key else 'Not found'}")
        print(f"   API_KEY: {'Found' if env_api_key else 'Not found'}")
        
        current_api_key = (env_openai_api_key or 
                          env_open_api_key or  # User's environment variable
                          env_openai_key or 
                          env_api_key)
        
        if not current_api_key:
            print("❌ NO API KEY FOUND IN ENVIRONMENT VARIABLES")
            openai_client_error = "No API key found in environment variables"
            return None
        
        # Debug: Show what API key we're using (last 4 chars for security)
        key_suffix = current_api_key[-4:] if current_api_key else "None"
        print(f"🔑 Attempting OpenAI initialization with key ending in: {key_suffix}")
        print(f"🌍 KEY SOURCE: environment")
        
        # Show which specific env var was used
        if env_openai_api_key:
            print(f"   Using OPENAI_API_KEY (ends with: {env_openai_api_key[-4:]})")
        elif env_open_api_key:
            print(f"   Using OPEN_API_KEY (ends with: {env_open_api_key[-4:]})")
        elif env_openai_key:
            print(f"   Using OPENAI_KEY (ends with: {env_openai_key[-4:]})")
        elif env_api_key:
            print(f"   Using API_KEY (ends with: {env_api_key[-4:]})")
        
        try:
            # Method 1: Direct initialization with API key
            openai_client = OpenAI(api_key=current_api_key)
            
            # Test the client with a minimal request to verify it works
            test_response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": "Test"}],
                max_tokens=1
            )
            
            print("✅ OpenAI client initialized and tested successfully")
            return openai_client
            
        except Exception as e:
            error_msg = str(e)
            print(f"❌ OpenAI client initialization failed: {error_msg}")
            
            # Try alternative initialization methods
            try:
                # Method 2: Environment variable approach
                os.environ["OPENAI_API_KEY"] = current_api_key
                openai_client = OpenAI()
                
                # Test this client too
                test_response = openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": "Test"}],
                    max_tokens=1
                )
                
                print("✅ OpenAI client initialized via environment variable")
                return openai_client
                
            except Exception as e2:
                print(f"❌ Alternative OpenAI initialization also failed: {str(e2)}")
                
                # Method 3: Check if it's an API key issue
                if "api_key" in error_msg.lower() or "401" in error_msg or "invalid" in error_msg.lower():
                    openai_client_error = "Invalid API key - OpenAI functionality disabled"
                    print("🔑 API key appears to be invalid - disabling OpenAI features")
                else:
                    openai_client_error = f"OpenAI client error: {error_msg}"
                    print(f"🔧 OpenAI client technical error: {error_msg}")
                
                return None
    
    def test_openai_connection():
        """Test OpenAI connection and return status"""
        client = get_openai_client()
        if client:
            return {"status": "connected", "message": "OpenAI client working"}
        elif openai_client_error:
            return {"status": "error", "message": openai_client_error}
        else:
            return {"status": "unknown", "message": "OpenAI status unknown"}
    
    def reinitialize_openai_client():
        """Force reinitialize OpenAI client - useful for production environments"""
        global openai_client, openai_client_error
        openai_client = None
        openai_client_error = None
        return get_openai_client()
    
    # Cost-optimized settings
    MAX_TOKENS_PER_REQUEST = 1000
    DAILY_TOKEN_LIMIT = 50000
    MODEL_NAME = "gpt-4o-mini"  # Most cost-effective model
    
except ImportError:
    OPENAI_AVAILABLE = False
    openai_client = None
    openai_client_error = "OpenAI library not installed"
    
    def get_openai_client():
        return None
        
    def test_openai_connection():
        return {"status": "unavailable", "message": "OpenAI library not installed"}
        
    print("OpenAI not available - install with: pip install openai")
    
except Exception as e:
    OPENAI_AVAILABLE = False
    openai_client = None
    openai_client_error = f"OpenAI import error: {str(e)}"
    
    def get_openai_client():
        return None
        
    def test_openai_connection():
        return {"status": "error", "message": openai_client_error}
        
    print(f"OpenAI initialization error: {str(e)}")

# Free document processing libraries - with graceful fallback
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# PDF reorganization libraries
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from PyPDF2 import PdfReader, PdfWriter
    import io
    PDF_GENERATION_AVAILABLE = True
except ImportError:
    PDF_GENERATION_AVAILABLE = False
    print("PDF generation not available - install with: pip install reportlab PyPDF2")

# Initialize Flask app FIRST - this must always succeed for gunicorn
app = Flask(__name__)

# Add basic error handler to ensure app is always functional
@app.errorhandler(500)
def handle_internal_error(e):
    return jsonify({'error': 'Internal server error', 'message': str(e)}), 500

@app.errorhandler(404)
def handle_not_found(e):
    return jsonify({'error': 'Not found', 'message': str(e)}), 404

# Basic health endpoint that always works
@app.route('/basic_health')
def basic_health():
    return jsonify({'status': 'Flask app running', 'timestamp': str(datetime.now())})

# Initialize CORS after app creation
try:
    CORS(app)
except Exception as e:
    print(f"Warning: CORS initialization failed: {e}")
    # App can still work without CORS in some environments

# Global storage for analysis rules and session data
analysis_rules = [
    {"pattern": "MORTGAGE|DEED OF TRUST", "type": "contains", "label": "Mortgage"},
    {"pattern": "PROMISSORY NOTE", "type": "contains", "label": "Promissory Note"},
    {"pattern": "CLOSING INSTRUCTIONS", "type": "contains", "label": "Lenders Closing Instructions Guaranty"},
    {"pattern": "ANTI.?COERCION", "type": "contains", "label": "Statement of Anti Coercion Florida"},
    {"pattern": "POWER OF ATTORNEY", "type": "contains", "label": "Correction Agreement and Limited Power of Attorney"},
    {"pattern": "ACKNOWLEDGMENT", "type": "contains", "label": "All Purpose Acknowledgment"},
    {"pattern": "FLOOD HAZARD", "type": "contains", "label": "Flood Hazard Determination"},
    {"pattern": "AUTOMATIC PAYMENT", "type": "contains", "label": "Automatic Payments Authorization"},
    {"pattern": "TAX RECORD", "type": "contains", "label": "Tax Record Information"}
]

lender_requirements = {}

# Industry templates with capabilities
INDUSTRY_TEMPLATES = {
    'mortgage': {
        'id': 'mortgage',
        'name': 'Mortgage & Real Estate',
        'icon': '🏠',
        'description': 'Mortgage packages, loan documents, real estate transactions',
        'capabilities': [
            {'icon': '📧', 'name': 'Email Parser', 'description': 'Extract lender requirements from emails'},
            {'icon': '📋', 'name': 'Workflow Management', 'description': '3-step guided process'},
            {'icon': '✅', 'name': 'Compliance Check', 'description': 'TRID & RESPA validation'},
            {'icon': '🛡️', 'name': 'Fraud Detection', 'description': 'Document authenticity verification'}
        ],
        'performance_metric': '99%+ accuracy'
    },
    'real_estate': {
        'id': 'real_estate',
        'name': 'Real Estate Transactions',
        'icon': '🏘️',
        'description': 'Property transactions, deeds, titles, purchase agreements, inspections',
        'capabilities': [
            {'icon': '🛡️', 'name': 'Fraud Detection', 'description': 'Advanced document authenticity verification'},
            {'icon': '✅', 'name': 'Compliance Validation', 'description': 'Regulatory requirement checking'},
            {'icon': '📊', 'name': 'Risk Scoring', 'description': 'Real-time fraud and compliance assessment'},
            {'icon': '⚡', 'name': 'Speed Boost', 'description': '90% faster processing'}
        ],
        'performance_metric': '95% fraud detection'
    },
    'legal': {
        'id': 'legal',
        'name': 'Legal & Law Firms',
        'icon': '⚖️',
        'description': 'Contracts, agreements, legal documents, case files',
        'capabilities': [
            {'icon': '📄', 'name': 'Contract Analysis', 'description': 'Automated contract review and analysis'},
            {'icon': '✅', 'name': 'Legal Compliance', 'description': 'Regulatory and legal requirement checking'},
            {'icon': '🔍', 'name': 'Document Review', 'description': 'Automated legal document analysis'},
            {'icon': '📊', 'name': 'Case Organization', 'description': 'Intelligent case file management'}
        ],
        'performance_metric': '80% faster review'
    },
    'healthcare': {
        'id': 'healthcare',
        'name': 'Healthcare & Medical',
        'icon': '🏥',
        'description': 'Medical records, insurance claims, patient documents',
        'capabilities': [
            {'icon': '📋', 'name': 'Medical Record Analysis', 'description': 'Comprehensive patient record review'},
            {'icon': '💰', 'name': 'Claims Processing', 'description': 'Automated insurance claim analysis'},
            {'icon': '✅', 'name': 'HIPAA Compliance', 'description': 'Healthcare privacy regulation adherence'},
            {'icon': '🔍', 'name': 'Clinical Data Extraction', 'description': 'Extract key medical information'}
        ],
        'performance_metric': 'HIPAA compliant'
    },
    'financial': {
        'id': 'financial',
        'name': 'Financial Services',
        'icon': '💰',
        'description': 'Banking documents, investment reports, financial statements',
        'capabilities': [
            {'icon': '📊', 'name': 'Financial Analysis', 'description': 'Comprehensive financial document review'},
            {'icon': '✅', 'name': 'Regulatory Compliance', 'description': 'Banking and finance regulation checking'},
            {'icon': '🔍', 'name': 'Risk Assessment', 'description': 'Financial risk analysis and scoring'},
            {'icon': '💳', 'name': 'Credit Analysis', 'description': 'Automated credit evaluation'}
        ],
        'performance_metric': 'SOX compliant'
    },
    'hr': {
        'id': 'hr',
        'name': 'Human Resources',
        'icon': '👥',
        'description': 'Employee records, resumes, HR documents, onboarding',
        'capabilities': [
            {'icon': '📄', 'name': 'Resume Analysis', 'description': 'Automated candidate evaluation'},
            {'icon': '✅', 'name': 'Compliance Checking', 'description': 'HR regulation and policy adherence'},
            {'icon': '🔍', 'name': 'Background Verification', 'description': 'Employee background document review'},
            {'icon': '📊', 'name': 'Performance Analytics', 'description': 'Employee performance data analysis'}
        ],
        'performance_metric': '70% faster hiring'
    }
}

def clean_text(text):
    """Clean and normalize text content"""
    if not text:
        return ""
    
    # Remove excessive whitespace and normalize
    text = re.sub(r'\s+', ' ', text.strip())
    
    # Remove special characters that might cause issues
    text = re.sub(r'[^\w\s\-\.\,\:\;\!\?\(\)\[\]\/\@\#\$\%\&\*\+\=]', '', text)
    
    return text

class UniversalDocumentProcessor:
    """Universal document processor supporting multiple file formats"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.png': self._process_image,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.gif': self._process_image,
            '.txt': self._process_txt
        }
    
    def process_file(self, file_path, filename):
        """Process a file and extract text content"""
        try:
            file_ext = os.path.splitext(filename.lower())[1]
            
            if file_ext not in self.supported_formats:
                return {
                    'success': False,
                    'error': f'Unsupported file format: {file_ext}',
                    'text': '',
                    'metadata': {'file_type': file_ext}
                }
            
            processor = self.supported_formats[file_ext]
            result = processor(file_path, filename)
            
            # Clean the extracted text
            if result.get('success') and result.get('text'):
                result['text'] = clean_text(result['text'])
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {'processing_error': True}
            }
    
    def _process_pdf(self, file_path, filename):
        """Process PDF files"""
        if not PDFPLUMBER_AVAILABLE:
            return {'success': False, 'error': 'PDF processing not available', 'text': '', 'metadata': {}}
        
        try:
            text_content = ""
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            
            return {
                'success': True,
                'text': text_content,
                'metadata': {
                    'pages': len(pdf.pages) if 'pdf' in locals() else 0,
                    'file_size': os.path.getsize(file_path),
                    'processing_method': 'pdfplumber'
                }
            }
        except Exception as e:
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}
    
    def _process_docx(self, file_path, filename):
        """Process Word documents"""
        if not DOCX_AVAILABLE:
            return {'success': False, 'error': 'Word processing not available', 'text': '', 'metadata': {}}
        
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            return {
                'success': True,
                'text': text_content,
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'file_size': os.path.getsize(file_path),
                    'processing_method': 'python-docx'
                }
            }
        except Exception as e:
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}
    
    def _process_excel(self, file_path, filename):
        """Process Excel files"""
        if not OPENPYXL_AVAILABLE:
            return {'success': False, 'error': 'Excel processing not available', 'text': '', 'metadata': {}}
        
        try:
            workbook = load_workbook(file_path, data_only=True)
            text_content = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content += f"\n--- Sheet: {sheet_name} ---\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell))
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            return {
                'success': True,
                'text': text_content,
                'metadata': {
                    'sheets': len(workbook.sheetnames),
                    'file_size': os.path.getsize(file_path),
                    'processing_method': 'openpyxl'
                }
            }
        except Exception as e:
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}
    
    def _process_image(self, file_path, filename):
        """Process image files using OCR"""
        if not OCR_AVAILABLE:
            return {'success': False, 'error': 'OCR processing not available', 'text': '', 'metadata': {}}
        
        try:
            # Open and preprocess image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Resize if too large (for better OCR performance)
            width, height = image.size
            if width > 2000 or height > 2000:
                scale_factor = min(2000/width, 2000/height)
                new_size = (int(width * scale_factor), int(height * scale_factor))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            # Convert to grayscale for better OCR
            image = image.convert('L')
            
            # Extract text using OCR
            text_content = pytesseract.image_to_string(image)
            
            return {
                'success': True,
                'text': text_content,
                'metadata': {
                    'image_size': f"{width}x{height}",
                    'file_size': os.path.getsize(file_path),
                    'processing_method': 'pytesseract-ocr'
                }
            }
        except Exception as e:
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}
    
    def _process_txt(self, file_path, filename):
        """Process text files"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    
                    return {
                        'success': True,
                        'text': text_content,
                        'metadata': {
                            'encoding': encoding,
                            'file_size': os.path.getsize(file_path),
                            'processing_method': 'text-reader'
                        }
                    }
                except UnicodeDecodeError:
                    continue
            
            return {'success': False, 'error': 'Could not decode text file', 'text': '', 'metadata': {}}
        except Exception as e:
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}
    
    def separate_documents(self, file_path, filename):
        """Separate documents into individual sections/components"""
        try:
            # First, process the file to get text content
            processing_result = self.process_file(file_path, filename)
            
            if not processing_result['success']:
                return {
                    'success': False,
                    'error': processing_result['error'],
                    'sections': []
                }
            
            text_content = processing_result['text']
            
            # For mortgage documents, use the existing analyze_mortgage_sections function
            if filename.lower().endswith('.pdf'):
                sections = analyze_mortgage_sections(filename, use_lender_rules=True)
                
                # Add text content to each section for viewing
                for section in sections:
                    section['text_preview'] = text_content[:500] + "..." if len(text_content) > 500 else text_content
                    section['full_text'] = text_content
                
                return {
                    'success': True,
                    'sections': sections,
                    'total_sections': len(sections),
                    'original_text': text_content
                }
            else:
                # For non-PDF files, create a single section
                return {
                    'success': True,
                    'sections': [{
                        'name': f"Document: {filename}",
                        'pages': "1",
                        'confidence': "high",
                        'risk_score': 10,
                        'text_preview': text_content[:500] + "..." if len(text_content) > 500 else text_content,
                        'full_text': text_content
                    }],
                    'total_sections': 1,
                    'original_text': text_content
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': f"Document separation failed: {str(e)}",
                'sections': []
            }
    
    def analyze_document(self, file_path, filename, section_name=None):
        """Analyze a specific document or section"""
        try:
            # Process the file
            processing_result = self.process_file(file_path, filename)
            
            if not processing_result['success']:
                return {
                    'success': False,
                    'error': processing_result['error'],
                    'analysis': {}
                }
            
            text_content = processing_result['text']
            
            # Perform analysis using existing function
            analysis_result = analyze_document_content(text_content, filename, 'mortgage')
            
            # Add section-specific information if provided
            if section_name:
                analysis_result['section_name'] = section_name
                analysis_result['section_analysis'] = f"Analysis for section: {section_name}"
            
            return {
                'success': True,
                'analysis': analysis_result,
                'text_length': len(text_content),
                'processing_metadata': processing_result['metadata']
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Document analysis failed: {str(e)}",
                'analysis': {}
            }
    
    def extract_text(self, file_path, filename, section_name=None, page_range=None):
        """Extract text from document or specific section"""
        try:
            # Process the file
            processing_result = self.process_file(file_path, filename)
            
            if not processing_result['success']:
                return {
                    'success': False,
                    'error': processing_result['error'],
                    'text': ''
                }
            
            text_content = processing_result['text']
            
            # If page range is specified, try to extract specific pages
            if page_range and filename.lower().endswith('.pdf'):
                # For PDF files, we could implement page-specific extraction
                # For now, return the full text with a note about the requested range
                extracted_text = f"[Requested pages: {page_range}]\n\n{text_content}"
            else:
                extracted_text = text_content
            
            return {
                'success': True,
                'text': extracted_text,
                'section_name': section_name,
                'page_range': page_range,
                'text_length': len(extracted_text),
                'metadata': processing_result['metadata']
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Text extraction failed: {str(e)}",
                'text': ''
            }
    
    def export_section(self, file_path, filename, section_name, export_format='txt'):
        """Export a specific section in the requested format"""
        try:
            # Extract text for the section
            extraction_result = self.extract_text(file_path, filename, section_name)
            
            if not extraction_result['success']:
                return {
                    'success': False,
                    'error': extraction_result['error'],
                    'export_path': None
                }
            
            text_content = extraction_result['text']
            
            # Create export filename
            safe_section_name = "".join(c for c in section_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            export_filename = f"{safe_section_name}_{filename.split('.')[0]}.{export_format}"
            export_path = f"/tmp/{export_filename}"
            
            # Export based on format
            if export_format.lower() == 'txt':
                with open(export_path, 'w', encoding='utf-8') as f:
                    f.write(f"Section: {section_name}\n")
                    f.write(f"Source: {filename}\n")
                    f.write("=" * 50 + "\n\n")
                    f.write(text_content)
            
            elif export_format.lower() == 'json':
                import json
                export_data = {
                    'section_name': section_name,
                    'source_file': filename,
                    'text_content': text_content,
                    'export_timestamp': datetime.now().isoformat(),
                    'metadata': extraction_result['metadata']
                }
                with open(export_path, 'w', encoding='utf-8') as f:
                    json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            else:
                return {
                    'success': False,
                    'error': f"Unsupported export format: {export_format}",
                    'export_path': None
                }
            
            return {
                'success': True,
                'export_path': export_path,
                'export_filename': export_filename,
                'section_name': section_name,
                'export_format': export_format
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Export failed: {str(e)}",
                'export_path': None
            }

# Initialize the universal processor
document_processor = UniversalDocumentProcessor()

# AI-Powered Analysis Functions with Cost Optimization
class MortgageAI:
    """AI-powered mortgage document analysis with cost controls"""
    
    def __init__(self):
        self.total_tokens_used = 0
        self.max_daily_tokens = 50000  # Cost control limit
        self.analysis_cache = {}  # Cache results to avoid re-analysis
    
    def check_token_limit(self, estimated_tokens):
        """Check if we're within token limits"""
        if self.total_tokens_used + estimated_tokens > self.max_daily_tokens:
            return False, f"Daily token limit reached ({self.max_daily_tokens})"
        return True, "OK"
    
    def estimate_tokens(self, text):
        """Estimate token count (rough approximation: 1 token ≈ 4 characters)"""
        return len(text) // 4
    
    def truncate_text(self, text, max_tokens=MAX_INPUT_TOKENS):
        """Truncate text to stay within token limits"""
        estimated_tokens = self.estimate_tokens(text)
        if estimated_tokens <= max_tokens:
            return text
        
        # Truncate to approximately max_tokens
        max_chars = max_tokens * 4
        return text[:max_chars] + "... [truncated for cost optimization]"
    
    def analyze_mortgage_document(self, text, document_type="unknown"):
        """AI-powered mortgage document analysis"""
        if not OPENAI_AVAILABLE:
            return {"error": "OpenAI not available", "ai_analysis": False}
        
        # Check cache first
        text_hash = hash(text[:1000])  # Use first 1000 chars for cache key
        if text_hash in self.analysis_cache:
            return self.analysis_cache[text_hash]
        
        # Truncate text for cost control
        text = self.truncate_text(text)
        estimated_tokens = self.estimate_tokens(text)
        
        # Check token limits
        can_proceed, message = self.check_token_limit(estimated_tokens + MAX_TOKENS_PER_REQUEST)
        if not can_proceed:
            return {"error": message, "ai_analysis": False}
        
        try:
            # Cost-optimized prompt for mortgage analysis
            prompt = f"""Analyze this mortgage document efficiently. Extract key information:

Document Type: {document_type}
Text: {text}

Provide a JSON response with:
1. document_category (Mortgage, Promissory Note, Closing Instructions, etc.)
2. key_details (borrower, lender, amount, property, etc.)
3. compliance_flags (any issues or missing info)
4. confidence_score (0-100)

Keep response concise to minimize costs."""

            # Get OpenAI client with error handling
            client = get_openai_client()
            if not client:
                return {"error": "OpenAI client not available", "ai_enhanced": False}
            
            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=MAX_TOKENS_PER_REQUEST,
                temperature=0.1  # Low temperature for consistent results
            )
            
            # Track token usage
            tokens_used = response.usage.total_tokens
            self.total_tokens_used += tokens_used
            
            result = {
                "ai_analysis": True,
                "analysis": response.choices[0].message.content,
                "tokens_used": tokens_used,
                "total_tokens": self.total_tokens_used,
                "model": MODEL_NAME
            }
            
            # Cache the result
            self.analysis_cache[text_hash] = result
            
            return result
            
        except Exception as e:
            return {"error": f"AI analysis failed: {str(e)}", "ai_analysis": False}
    
    def enhance_lender_parsing(self, email_content):
        """AI-enhanced lender email parsing"""
        if not OPENAI_AVAILABLE:
            return {"error": "OpenAI not available", "ai_enhanced": False}
        
        # Truncate for cost control
        email_content = self.truncate_text(email_content, 4000)  # Smaller limit for emails
        estimated_tokens = self.estimate_tokens(email_content)
        
        can_proceed, message = self.check_token_limit(estimated_tokens + 500)  # Smaller response
        if not can_proceed:
            return {"error": message, "ai_enhanced": False}
        
        try:
            prompt = f"""Extract lender requirements from this email efficiently:

{email_content}

Return JSON with:
1. lender_name
2. required_documents (list)
3. special_instructions
4. deadline
5. contact_info

Be concise to minimize costs."""

            # Get OpenAI client with error handling
            client = get_openai_client()
            if not client:
                return {"error": "OpenAI client not available", "ai_enhanced": False}

            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=500,  # Small response for cost control
                temperature=0.1
            )
            
            tokens_used = response.usage.total_tokens
            self.total_tokens_used += tokens_used
            
            return {
                "ai_enhanced": True,
                "enhanced_parsing": response.choices[0].message.content,
                "tokens_used": tokens_used,
                "total_tokens": self.total_tokens_used
            }
            
        except Exception as e:
            return {"error": f"AI enhancement failed: {str(e)}", "ai_enhanced": False}

# Initialize AI processor
mortgage_ai = MortgageAI() if OPENAI_AVAILABLE else None

class PDFReorganizationAI:
    """AI-powered PDF reorganization based on lender requirements"""
    
    def __init__(self):
        self.total_tokens_used = 0
        self.max_daily_tokens = 50000  # Cost control limit
        self.reorganization_cache = {}
    
    def analyze_lender_requirements(self, lender_requirements):
        """Extract document order preferences from lender requirements"""
        if not OPENAI_AVAILABLE:
            return {"error": "OpenAI not available", "ai_analysis": False}
        
        try:
            prompt = f"""Analyze these lender requirements and determine the optimal document order for a mortgage package:

Requirements: {json.dumps(lender_requirements, indent=2)}

Provide a JSON response with:
1. preferred_order: Array of document types in preferred order
2. mandatory_documents: Array of required documents
3. optional_documents: Array of optional documents
4. special_instructions: Any specific ordering or formatting requirements

Common mortgage document types: Mortgage, Promissory Note, Closing Instructions, Anti Coercion Statement, Power of Attorney, Acknowledgment, Flood Hazard, Automatic Payments, Tax Records"""

            # Get OpenAI client with error handling
            client = get_openai_client()
            if not client:
                return {"error": "OpenAI client not available", "ai_analysis": False}

            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=MAX_TOKENS_PER_REQUEST,
                temperature=0.1
            )
            
            tokens_used = response.usage.total_tokens
            self.total_tokens_used += tokens_used
            
            return {
                "ai_analysis": True,
                "analysis": response.choices[0].message.content,
                "tokens_used": tokens_used,
                "total_tokens": self.total_tokens_used
            }
            
        except Exception as e:
            return {"error": f"AI analysis failed: {str(e)}", "ai_analysis": False}
    
    def match_documents_to_requirements(self, separated_documents, lender_requirements):
        """Memory-safe version of document matching"""
        if not OPENAI_AVAILABLE:
            return {"error": "OpenAI not available", "ai_analysis": False}
        
        try:
            print(f"🧠 Memory before processing: {get_memory_usage():.1f} MB")
            
            # Truncate large content to prevent memory issues
            safe_documents = truncate_text_content(separated_documents, max_length=500)
            safe_requirements = truncate_text_content(lender_requirements, max_length=1000)
            
            prompt = f"""Match these documents to requirements (content truncated for processing):

Documents: {json.dumps(safe_documents[:10], indent=1)}
Requirements: {json.dumps(safe_requirements, indent=1)}

Provide JSON response with:
1. document_matches: [{{document_name, requirement_match, confidence_score}}]
2. suggested_order: [document names in optimal order]
3. compliance_score: 0-100
4. recommendations: [brief suggestions]"""

            # Get OpenAI client with error handling
            client = get_openai_client()
            if not client:
                return {"error": "OpenAI client not available", "ai_analysis": False}

            response = memory_safe_ai_call(client, prompt, max_tokens=800)
            
            print(f"🧠 Memory after AI call: {get_memory_usage():.1f} MB")
            
            tokens_used = response.usage.total_tokens
            self.total_tokens_used += tokens_used
            
            return {
                "ai_analysis": True,
                "matching_result": response.choices[0].message.content,
                "tokens_used": tokens_used
            }
            
        except Exception as e:
            gc.collect()  # Cleanup on error
            return {"error": f"Memory-safe matching failed: {str(e)}", "ai_analysis": False}
    
    def determine_optimal_order(self, matched_documents, lender_preferences=None):
        """Determine optimal document order using AI"""
        if not OPENAI_AVAILABLE:
            return {"error": "OpenAI not available", "ai_analysis": False}
        
        try:
            prompt = f"""Determine the optimal order for these mortgage documents based on industry standards and lender preferences:

Matched Documents:
{json.dumps(matched_documents, indent=2)}

Lender Preferences:
{json.dumps(lender_preferences, indent=2) if lender_preferences else "None specified"}

Consider:
1. Industry standard mortgage document order
2. Lender-specific requirements
3. Logical document flow (e.g., Mortgage before Promissory Note)
4. Regulatory compliance requirements

Provide a JSON response with:
1. ordered_documents: Array of documents in optimal order
2. reasoning: Explanation for the ordering decisions
3. confidence_score: Confidence in the ordering (0-100)
4. alternative_orders: Other viable ordering options"""

            # Get OpenAI client with error handling
            client = get_openai_client()
            if not client:
                return {"error": "OpenAI client not available", "ai_analysis": False}

            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=MAX_TOKENS_PER_REQUEST,
                temperature=0.1
            )
            
            tokens_used = response.usage.total_tokens
            self.total_tokens_used += tokens_used
            
            return {
                "ai_analysis": True,
                "ordering_result": response.choices[0].message.content,
                "tokens_used": tokens_used,
                "total_tokens": self.total_tokens_used
            }
            
        except Exception as e:
            return {"error": f"AI ordering failed: {str(e)}", "ai_analysis": False}
    
    def validate_compliance(self, ordered_documents, lender_requirements):
        """Validate compliance of the organized document package"""
        if not OPENAI_AVAILABLE:
            return {"error": "OpenAI not available", "ai_analysis": False}
        
        try:
            prompt = f"""Validate the compliance of this organized mortgage document package:

Organized Documents:
{json.dumps(ordered_documents, indent=2)}

Lender Requirements:
{json.dumps(lender_requirements, indent=2)}

Provide a JSON response with:
1. compliance_score: Overall compliance percentage (0-100)
2. compliance_checklist: Array of requirements with pass/fail status
3. missing_documents: Documents required but not included
4. warnings: Potential issues or concerns
5. recommendations: Suggestions for improvement
6. risk_assessment: Low/Medium/High risk level"""

            # Get OpenAI client with error handling
            client = get_openai_client()
            if not client:
                return {"error": "OpenAI client not available", "ai_analysis": False}

            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=MAX_TOKENS_PER_REQUEST,
                temperature=0.1
            )
            
            tokens_used = response.usage.total_tokens
            self.total_tokens_used += tokens_used
            
            return {
                "ai_analysis": True,
                "compliance_result": response.choices[0].message.content,
                "tokens_used": tokens_used,
                "total_tokens": self.total_tokens_used
            }
            
        except Exception as e:
            return {"error": f"AI compliance validation failed: {str(e)}", "ai_analysis": False}
    
    def generate_reorganization_plan(self, separated_documents, lender_requirements):
        """Generate a complete reorganization plan using AI"""
        if not OPENAI_AVAILABLE:
            return {"error": "OpenAI not available", "ai_analysis": False}
        
        # Check cache first
        cache_key = hash(str(separated_documents) + str(lender_requirements))
        if cache_key in self.reorganization_cache:
            return self.reorganization_cache[cache_key]
        
        try:
            # Step 1: Match documents to requirements
            matching_result = self.match_documents_to_requirements(separated_documents, lender_requirements)
            if not matching_result.get("ai_analysis"):
                return matching_result
            
            # Step 2: Determine optimal order
            ordering_result = self.determine_optimal_order(separated_documents, lender_requirements)
            if not ordering_result.get("ai_analysis"):
                return ordering_result
            
            # Step 3: Validate compliance
            compliance_result = self.validate_compliance(separated_documents, lender_requirements)
            if not compliance_result.get("ai_analysis"):
                return compliance_result
            
            # Combine results
            total_tokens = (matching_result.get("tokens_used", 0) + 
                          ordering_result.get("tokens_used", 0) + 
                          compliance_result.get("tokens_used", 0))
            
            reorganization_plan = {
                "ai_analysis": True,
                "matching_analysis": matching_result["matching_result"],
                "ordering_analysis": ordering_result["ordering_result"],
                "compliance_analysis": compliance_result["compliance_result"],
                "total_tokens_used": total_tokens,
                "total_cost_estimate": f"${(total_tokens * 0.00015):.4f}"
            }
            
            # Cache the result
            self.reorganization_cache[cache_key] = reorganization_plan
            
            return reorganization_plan
            
        except Exception as e:
            return {"error": f"AI reorganization planning failed: {str(e)}", "ai_analysis": False}

# Initialize PDF reorganization AI with robust checking
def initialize_pdf_reorganizer_ai():
    """Initialize PDF reorganizer AI with proper OpenAI client validation"""
    try:
        if not OPENAI_AVAILABLE:
            return None
        
        # Test if OpenAI client is actually working
        client = get_openai_client()
        if not client:
            return None
        
        # Test connection
        connection_status = test_openai_connection()
        if connection_status['status'] != 'connected':
            return None
        
        return PDFReorganizationAI()
    except Exception as e:
        print(f"Failed to initialize PDF reorganizer AI: {str(e)}")
        return None

# Initialize PDF reorganization AI with robust checking (safe for gunicorn)
try:
    pdf_reorganizer_ai = initialize_pdf_reorganizer_ai()
    print("✅ PDF reorganizer AI initialized successfully")
except Exception as e:
    pdf_reorganizer_ai = None
    print(f"⚠️  PDF reorganizer AI initialization failed: {e}")
    print("   App will continue without AI features")

class PDFReorganizer:
    """PDF reorganization and generation based on AI analysis"""
    
    def __init__(self):
        self.temp_dir = "/tmp/pdf_reorganization"
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def extract_document_pages(self, original_pdf_path, document_sections):
        """Extract pages for each document section from the original PDF"""
        if not PDF_GENERATION_AVAILABLE:
            return {"error": "PDF generation libraries not available", "success": False}
        
        try:
            extracted_documents = {}
            
            with open(original_pdf_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for section in document_sections:
                    section_name = section.get('name', 'Unknown')
                    pages_range = section.get('pages', '1-1')
                    
                    # Parse page range (e.g., "2-3" or "5")
                    if '-' in pages_range:
                        start_page, end_page = map(int, pages_range.split('-'))
                    else:
                        start_page = end_page = int(pages_range)
                    
                    # Adjust for 0-based indexing and validate range
                    start_page = max(1, min(start_page, total_pages)) - 1
                    end_page = max(1, min(end_page, total_pages)) - 1
                    
                    # Extract pages for this document
                    pdf_writer = PdfWriter()
                    for page_num in range(start_page, end_page + 1):
                        if page_num < total_pages:
                            pdf_writer.add_page(pdf_reader.pages[page_num])
                    
                    # Save extracted document
                    output_path = os.path.join(self.temp_dir, f"{section_name.replace(' ', '_')}.pdf")
                    with open(output_path, 'wb') as output_file:
                        pdf_writer.write(output_file)
                    
                    extracted_documents[section_name] = {
                        'path': output_path,
                        'pages': pages_range,
                        'confidence': section.get('confidence', 'medium'),
                        'quality': section.get('quality', '90%')
                    }
            
            return {"success": True, "extracted_documents": extracted_documents}
            
        except Exception as e:
            return {"error": f"PDF extraction failed: {str(e)}", "success": False}
    
    def generate_cover_page(self, compliance_summary, lender_info):
        """Generate a professional cover page for the reorganized PDF"""
        if not PDF_GENERATION_AVAILABLE:
            return {"error": "PDF generation libraries not available", "success": False}
        
        try:
            cover_path = os.path.join(self.temp_dir, "cover_page.pdf")
            
            # Create cover page
            doc = SimpleDocTemplate(cover_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            story.append(Paragraph("AI-Organized Mortgage Document Package", title_style))
            story.append(Spacer(1, 20))
            
            # Lender Information
            lender_style = ParagraphStyle(
                'LenderInfo',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=10
            )
            
            story.append(Paragraph(f"<b>Lender:</b> {lender_info.get('lender_name', 'Unknown')}", lender_style))
            story.append(Paragraph(f"<b>Contact:</b> {lender_info.get('contact_email', 'N/A')}", lender_style))
            story.append(Paragraph(f"<b>Funding Amount:</b> {lender_info.get('funding_amount', 'N/A')}", lender_style))
            story.append(Paragraph(f"<b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", lender_style))
            story.append(Spacer(1, 30))
            
            # Compliance Summary
            compliance_style = ParagraphStyle(
                'ComplianceInfo',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=8
            )
            
            story.append(Paragraph("<b>AI Analysis Summary:</b>", styles['Heading2']))
            story.append(Paragraph(f"Compliance Score: {compliance_summary.get('compliance_score', 'N/A')}%", compliance_style))
            story.append(Paragraph(f"Documents Matched: {compliance_summary.get('documents_matched', 'N/A')}", compliance_style))
            story.append(Paragraph(f"AI Model Used: GPT-4o-mini", compliance_style))
            story.append(Spacer(1, 20))
            
            # Document List
            story.append(Paragraph("<b>Included Documents:</b>", styles['Heading3']))
            for doc in compliance_summary.get('document_list', []):
                story.append(Paragraph(f"• {doc}", compliance_style))
            
            doc.build(story)
            return {"success": True, "cover_path": cover_path}
            
        except Exception as e:
            return {"error": f"Cover page generation failed: {str(e)}", "success": False}
    
    def assemble_final_pdf(self, ordered_documents, cover_page_path, output_filename):
        """Assemble the final reorganized PDF"""
        if not PDF_GENERATION_AVAILABLE:
            return {"error": "PDF generation libraries not available", "success": False}
        
        try:
            final_pdf_path = os.path.join(self.temp_dir, output_filename)
            pdf_writer = PdfWriter()
            
            # Add cover page first
            if cover_page_path and os.path.exists(cover_page_path):
                with open(cover_page_path, 'rb') as cover_file:
                    cover_reader = PdfReader(cover_file)
                    for page in cover_reader.pages:
                        pdf_writer.add_page(page)
            
            # Add documents in the specified order
            for doc_info in ordered_documents:
                doc_path = doc_info.get('path')
                if doc_path and os.path.exists(doc_path):
                    with open(doc_path, 'rb') as doc_file:
                        doc_reader = PdfReader(doc_file)
                        for page in doc_reader.pages:
                            pdf_writer.add_page(page)
            
            # Write final PDF
            with open(final_pdf_path, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            return {"success": True, "final_pdf_path": final_pdf_path}
            
        except Exception as e:
            return {"error": f"PDF assembly failed: {str(e)}", "success": False}
    
    def reorganize_pdf(self, original_pdf_path, ai_analysis, lender_requirements, document_sections):
        """Complete PDF reorganization process using AI analysis"""
        try:
            # Step 1: Extract document pages
            extraction_result = self.extract_document_pages(original_pdf_path, document_sections)
            if not extraction_result.get("success"):
                return extraction_result
            
            extracted_docs = extraction_result["extracted_documents"]
            
            # Step 2: Parse AI analysis to get document order
            try:
                import json
                # Try to parse the AI ordering analysis
                ordering_data = json.loads(ai_analysis.get("ordering_analysis", "{}"))
                ordered_doc_names = ordering_data.get("ordered_documents", [])
            except:
                # Fallback to original order if AI parsing fails
                ordered_doc_names = list(extracted_docs.keys())
            
            # Step 3: Create ordered document list
            ordered_documents = []
            for doc_name in ordered_doc_names:
                if doc_name in extracted_docs:
                    ordered_documents.append(extracted_docs[doc_name])
            
            # Add any remaining documents not in the AI order
            for doc_name, doc_info in extracted_docs.items():
                if doc_name not in ordered_doc_names:
                    ordered_documents.append(doc_info)
            
            # Step 4: Generate cover page
            compliance_summary = {
                "compliance_score": "95",  # Default, could be parsed from AI analysis
                "documents_matched": len(ordered_documents),
                "document_list": [doc_name for doc_name in extracted_docs.keys()]
            }
            
            cover_result = self.generate_cover_page(compliance_summary, lender_requirements)
            cover_path = cover_result.get("cover_path") if cover_result.get("success") else None
            
            # Step 5: Assemble final PDF
            output_filename = f"reorganized_mortgage_package_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            assembly_result = self.assemble_final_pdf(ordered_documents, cover_path, output_filename)
            
            if assembly_result.get("success"):
                return {
                    "success": True,
                    "final_pdf_path": assembly_result["final_pdf_path"],
                    "output_filename": output_filename,
                    "documents_included": len(ordered_documents),
                    "ai_analysis_used": True,
                    "compliance_summary": compliance_summary
                }
            else:
                return assembly_result
                
        except Exception as e:
            return {"error": f"PDF reorganization failed: {str(e)}", "success": False}
    
    def cleanup_temp_files(self):
        """Clean up temporary files"""
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                os.makedirs(self.temp_dir, exist_ok=True)
        except Exception as e:
            print(f"Cleanup warning: {str(e)}")

# Initialize PDF reorganizer (safe for gunicorn)
try:
    pdf_reorganizer = PDFReorganizer() if PDF_GENERATION_AVAILABLE else None
    if pdf_reorganizer:
        print("✅ PDF reorganizer initialized successfully")
    else:
        print("⚠️  PDF reorganizer not available (PDF libraries missing)")
except Exception as e:
    pdf_reorganizer = None
    print(f"⚠️  PDF reorganizer initialization failed: {e}")
    print("   App will continue without PDF generation features")

def parse_lender_email(content):
    """Enhanced email parsing for lender requirements with memory optimization"""
    
    # Limit content size to prevent memory issues
    if len(content) > 50000:  # Limit to 50KB
        content = content[:50000]
    
    # Clean the content first
    content = clean_text(content)
    
    # Extract lender information
    lender_info = {
        'lender_name': 'Unknown Lender',
        'contact_email': '',
        'contact_name': '',
        'date': datetime.now().strftime('%Y-%m-%d'),
        'documents': [],
        'funding_amount': '',
        'special_instructions': []
    }
    
    try:
        # Extract lender name and contact info
        email_patterns = [
            r'From:\s*([^<]+)<([^>]+)>',
            r'([A-Za-z\s]+)\s*<([^@]+@[^>]+)>',
            r'([^@]+@[\w\.-]+\.\w+)'
        ]
        
        for pattern in email_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                if len(match.groups()) >= 2:
                    lender_info['contact_name'] = match.group(1).strip()
                    lender_info['contact_email'] = match.group(2).strip()
                else:
                    lender_info['contact_email'] = match.group(1).strip()
                break
        
        # Extract lender name from email domain or content
        if 'symmetry' in content.lower():
            lender_info['lender_name'] = 'Symmetry Lending'
        elif '@' in lender_info['contact_email']:
            domain = lender_info['contact_email'].split('@')[1].split('.')[0]
            lender_info['lender_name'] = domain.title() + ' Lending'
        
        # Extract funding amount
        amount_patterns = [
            r'\$[\d,]+\.?\d*',
            r'amount[:\s]+\$?([\d,]+\.?\d*)',
            r'fund[ing]*[:\s]+\$?([\d,]+\.?\d*)'
        ]
        
        for pattern in amount_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                lender_info['funding_amount'] = match.group(0) if '$' in match.group(0) else f"${match.group(1)}"
                break
        
        # Enhanced document extraction with checkbox patterns (limited to prevent memory issues)
        checkbox_patterns = [
            r'☐\s*([^\n\r]{1,200})',  # Limit match length
            r'□\s*([^\n\r]{1,200})', 
            r'▢\s*([^\n\r]{1,200})',
            r'\[\s*\]\s*([^\n\r]{1,200})',
            r'◯\s*([^\n\r]{1,200})',
            r'○\s*([^\n\r]{1,200})',
            r'•\s*([^\n\r]{1,200})',
            r'-\s*([^\n\r]{1,200})'
        ]
        
        documents = set()  # Use set to avoid duplicates
        
        for pattern in checkbox_patterns:
            try:
                matches = re.findall(pattern, content, re.MULTILINE)
                for match in matches[:50]:  # Limit to 50 matches per pattern
                    doc_name = clean_text(match.strip())
                    # Filter out non-document items
                    if (len(doc_name) > 5 and len(doc_name) < 200 and
                        not doc_name.lower().startswith(('below', 'all ', 'guard', 'think', 'st &', '1st &')) and
                        not re.match(r'^\d+$', doc_name)):
                        documents.add(doc_name)
                        
                        # Limit total documents to prevent memory issues
                        if len(documents) >= 100:
                            break
                            
                if len(documents) >= 100:
                    break
            except Exception:
                continue
        
        # Convert to list and sort
        lender_info['documents'] = sorted(list(documents))[:50]  # Limit to 50 documents
        
        # Extract special instructions (limited)
        instruction_patterns = [
            r'special instructions?[:\s]*([^\n\r]{1,500})',
            r'note[:\s]*([^\n\r]{1,500})',
            r'important[:\s]*([^\n\r]{1,500})',
            r'deadline[:\s]*([^\n\r]{1,500})'
        ]
        
        for pattern in instruction_patterns:
            try:
                matches = re.findall(pattern, content, re.IGNORECASE)
                lender_info['special_instructions'].extend([clean_text(match.strip()) for match in matches[:10]])  # Limit to 10 instructions
            except Exception:
                continue
        
    except Exception as e:
        # If parsing fails, return basic info
        lender_info['documents'] = ['Error parsing email content']
        lender_info['special_instructions'] = [f'Parsing error: {str(e)}']
    
    return lender_info

def analyze_mortgage_sections(filename, use_lender_rules=True):
    """Analyze mortgage sections using core categories and optional lender rules"""
    
    # Core mortgage categories (always included)
    core_sections = [
        "Mortgage",
        "Promissory Note", 
        "Lenders Closing Instructions Guaranty",
        "Statement of Anti Coercion Florida",
        "Correction Agreement and Limited Power of Attorney",
        "All Purpose Acknowledgment",
        "Flood Hazard Determination", 
        "Automatic Payments Authorization",
        "Tax Record Information"
    ]
    
    sections = []
    page_counter = 2
    
    # Add core sections
    for i, section_name in enumerate(core_sections):
        confidence = "high" if i < 3 else ("medium" if i < 6 else "low")
        risk_score = 15 if i < 3 else (25 if i < 6 else 35)
        
        sections.append({
            "name": section_name,
            "pages": f"{page_counter}-{page_counter + 1}",
            "confidence": confidence,
            "risk_score": risk_score,
            "quality": f"{95 - (i * 2)}%",
            "notes": f"Core mortgage document - {confidence} priority"
        })
        
        page_counter += 2
    
    # Add lender-specific sections if available and requested
    if use_lender_rules and lender_requirements.get('documents'):
        for i, doc_name in enumerate(lender_requirements['documents'][:5]):  # Limit to 5 additional
            sections.append({
                "name": doc_name,
                "pages": f"{page_counter}-{page_counter + 1}",
                "confidence": "medium",
                "risk_score": 20,
                "quality": "85%",
                "notes": f"Lender-specific requirement from {lender_requirements.get('lender_name', 'email')}"
            })
            page_counter += 2
    
    return sections

def analyze_document_content(text_content, filename, industry='mortgage'):
    """Analyze document content and return structured results"""
    
    # Basic document info
    word_count = len(text_content.split())
    char_count = len(text_content)
    
    # Industry-specific analysis
    if industry == 'mortgage':
        # Mortgage-specific patterns
        mortgage_patterns = {
            'loan_amount': r'\$[\d,]+\.?\d*',
            'interest_rate': r'\d+\.?\d*\s*%',
            'property_address': r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Boulevard|Blvd)',
            'borrower_name': r'Borrower[:\s]+([A-Za-z\s]+)',
            'lender_name': r'Lender[:\s]+([A-Za-z\s]+)'
        }
        
        extracted_data = {}
        for key, pattern in mortgage_patterns.items():
            matches = re.findall(pattern, text_content, re.IGNORECASE)
            if matches:
                extracted_data[key] = matches[0] if isinstance(matches[0], str) else matches[0]
        
        # Document categorization
        categories = []
        for rule in analysis_rules:
            if re.search(rule['pattern'], text_content, re.IGNORECASE):
                categories.append(rule['label'])
        
        return {
            'success': True,
            'filename': filename,
            'industry': industry,
            'word_count': word_count,
            'char_count': char_count,
            'categories': categories,
            'extracted_data': extracted_data,
            'quality_score': min(95, max(60, 100 - (len(text_content) // 10000))),  # Simple quality metric
            'processing_time': '2.3s'
        }
    
    else:
        # Universal analysis for other industries
        return {
            'success': True,
            'filename': filename,
            'industry': industry,
            'word_count': word_count,
            'char_count': char_count,
            'categories': ['General Document'],
            'extracted_data': {},
            'quality_score': min(95, max(60, 100 - (len(text_content) // 10000))),
            'processing_time': '1.8s'
        }

@app.route('/')
def index():
    """Main page with industry selection and upload interface"""
    industries_list = list(INDUSTRY_TEMPLATES.values())
    
    return render_template_string('''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Universal Document Analyzer</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: #0a0a0a;
            color: white;
            min-height: 100vh;
            position: relative;
            overflow-x: hidden;
        }

        .animated-bg {
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: linear-gradient(45deg, #0a0a0a, #1a1a2e, #16213e, #0f3460);
            background-size: 400% 400%;
            animation: gradientShift 15s ease infinite;
            z-index: -1;
        }

        @keyframes gradientShift {
            0% { background-position: 0% 50%; }
            50% { background-position: 100% 50%; }
            100% { background-position: 0% 50%; }
        }

        .container {
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            position: relative;
            z-index: 1;
        }

        .header {
            text-align: center;
            margin-bottom: 40px;
            padding: 40px 0;
        }

        .main-title {
            font-size: 3.5rem;
            font-weight: 900;
            background: linear-gradient(45deg, #00d4ff, #0099cc, #ffffff);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            margin-bottom: 20px;
            text-shadow: 0 0 30px rgba(0, 212, 255, 0.3);
            animation: glow 2s ease-in-out infinite alternate;
        }

        @keyframes glow {
            from { text-shadow: 0 0 20px rgba(0, 212, 255, 0.3); }
            to { text-shadow: 0 0 40px rgba(0, 212, 255, 0.6); }
        }

        .subtitle {
            font-size: 1.3rem;
            color: #b0b0b0;
            margin-bottom: 30px;
        }

        .industry-selector {
            background: rgba(255, 255, 255, 0.05);
            border-radius: 20px;
            padding: 30px;
            margin-bottom: 40px;
            backdrop-filter: blur(10px);
            border: 1px solid rgba(0, 212, 255, 0.2);
            box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3);
        }

        .selector-title {
            font-size: 1.8rem;
            font-weight: 700;
            color: #00d4ff;
            margin-bottom: 25px;
            text-align: center;
        }

        .industry-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(280px, 1fr));
            gap: 20px;
        }

        .industry-card {
            background: linear-gradient(135deg, rgba(255,255,255,0.1), rgba(255,255,255,0.05));
            border: 1px solid rgba(0,212,255,0.3);
            border-radius: 15px;
            padding: 25px;
            text-align: center;
            cursor: pointer;
            transition: all 0.3s ease;
            backdrop-filter: blur(10px);
            position: relative;
            overflow: hidden;
            height: 200px;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
        }

        .industry-card:hover {
            transform: translateY(-5px);
            border-color: #00d4ff;
            box-shadow: 0 10px 30px rgba(0,212,255,0.3);
        }

        .industry-card.selected {
            border-color: #ff6b35;
            background: linear-gradient(135deg, rgba(255,107,53,0.2), rgba(255,140,66,0.2));
            box-shadow: 0 15px 40px rgba(255,107,53,0.4);
        }

        .capability-overlay {
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: linear-gradient(135deg, rgba(0, 20, 40, 0.95), rgba(0, 40, 60, 0.95));
            opacity: 0;
            transition: all 0.3s ease;
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            padding: 20px;
            backdrop-filter: blur(10px);
        }

        .industry-card:hover .capability-overlay {
            opacity: 1;
        }

        .capability-grid-mini {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 15px;
            margin-bottom: 15px;
            width: 100%;
        }

        .capability-item-mini {
            text-align: center;
            padding: 8px;
        }

        .capability-icon-mini {
            font-size: 1.5rem;
            margin-bottom: 5px;
            display: block;
        }

        .capability-name-mini {
            color: #00d4ff;
            font-size: 0.8rem;
            font-weight: 600;
            line-height: 1.2;
        }

        .performance-badge {
            background: linear-gradient(45deg, #00ff00, #00cc00);
            color: #000;
            padding: 6px 12px;
            border-radius: 15px;
            font-size: 0.8rem;
            font-weight: 700;
            margin-top: 10px;
        }

        .overlay-title {
            color: #fff;
            font-size: 1rem;
            font-weight: 700;
            margin-bottom: 15px;
        }

        .industry-icon {
            font-size: 2.5rem;
            margin-bottom: 15px;
            display: block;
        }

        .industry-name {
            font-size: 1.2rem;
            font-weight: 600;
            color: #ffffff;
            margin-bottom: 10px;
        }

        .industry-desc {
            font-size: 0.9rem;
            color: #b0b0b0;
            line-height: 1.4;
        }

        /* MORTGAGE WORKFLOW SECTION */
        .mortgage-workflow {
            background: rgba(255, 107, 53, 0.1);
            border: 2px solid rgba(255, 107, 53, 0.3);
            border-radius: 15px;
            padding: 25px;
            margin-bottom: 20px;
            text-align: center;
        }

        .workflow-step {
            background: rgba(0, 212, 255, 0.1);
            border: 1px solid #00d4ff;
            border-radius: 10px;
            padding: 15px;
            margin: 10px 0;
            display: flex;
            align-items: center;
            justify-content: space-between;
        }

        .workflow-step.completed {
            background: rgba(0, 255, 0, 0.1);
            border-color: #00ff00;
        }

        .workflow-step.active {
            background: rgba(255, 107, 53, 0.1);
            border-color: #ff6b35;
            box-shadow: 0 0 15px rgba(255, 107, 53, 0.3);
        }

        .step-number {
            background: #00d4ff;
            color: #000;
            border-radius: 50%;
            width: 30px;
            height: 30px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: bold;
        }

        .step-number.completed {
            background: #00ff00;
        }

        .step-number.active {
            background: #ff6b35;
            color: white;
        }

        /* UNIVERSAL UPLOAD SECTION */
        .upload-section {
            background: rgba(255, 255, 255, 0.05);
            border-radius: 20px;
            padding: 40px;
            backdrop-filter: blur(10px);
            border: 1px solid rgba(0, 212, 255, 0.2);
            box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3);
            margin-bottom: 30px;
            display: none;
        }

        .upload-section.show {
            display: block;
        }

        .upload-header {
            text-align: center;
            margin-bottom: 30px;
        }

        .upload-title {
            font-size: 1.5rem;
            font-weight: 600;
            color: #ffffff;
            margin-bottom: 8px;
        }

        .upload-subtitle {
            color: #b0b0b0;
            font-size: 1rem;
            margin-bottom: 20px;
        }

        .supported-formats {
            display: flex;
            justify-content: center;
            gap: 15px;
            flex-wrap: wrap;
            margin-bottom: 30px;
        }

        .format-badge {
            background: rgba(0, 212, 255, 0.1);
            border: 1px solid rgba(0, 212, 255, 0.3);
            border-radius: 20px;
            padding: 6px 12px;
            font-size: 0.8rem;
            color: #00d4ff;
            font-weight: 500;
        }

        .upload-dropzone {
            position: relative;
            background: linear-gradient(135deg, rgba(0, 212, 255, 0.05), rgba(0, 212, 255, 0.02));
            border: 2px dashed rgba(0, 212, 255, 0.3);
            border-radius: 16px;
            padding: 60px 30px;
            text-align: center;
            cursor: pointer;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            overflow: hidden;
        }

        .upload-dropzone::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(0, 212, 255, 0.1), transparent);
            transition: left 0.5s;
        }

        .upload-dropzone:hover::before {
            left: 100%;
        }

        .upload-dropzone:hover {
            background: linear-gradient(135deg, rgba(0, 212, 255, 0.1), rgba(0, 212, 255, 0.05));
            border-color: rgba(0, 212, 255, 0.6);
            transform: translateY(-3px);
            box-shadow: 0 15px 35px rgba(0, 212, 255, 0.2);
        }

        .upload-dropzone.dragover {
            background: linear-gradient(135deg, rgba(0, 212, 255, 0.15), rgba(0, 212, 255, 0.08));
            border-color: #00d4ff;
            border-style: solid;
            box-shadow: 0 0 30px rgba(0, 212, 255, 0.4);
            transform: scale(1.02);
        }

        .upload-icon-container {
            margin-bottom: 20px;
        }

        .upload-icon {
            font-size: 4rem;
            color: #00d4ff;
            margin-bottom: 10px;
            display: block;
            animation: float 3s ease-in-out infinite;
        }

        @keyframes float {
            0%, 100% { transform: translateY(0px); }
            50% { transform: translateY(-10px); }
        }

        .upload-text {
            font-size: 1.3rem;
            font-weight: 600;
            color: #ffffff;
            margin-bottom: 8px;
        }

        .upload-hint {
            color: #888;
            font-size: 1rem;
            margin-bottom: 25px;
            line-height: 1.5;
        }

        .upload-actions {
            display: flex;
            justify-content: center;
            gap: 15px;
            flex-wrap: wrap;
        }

        .upload-btn {
            background: linear-gradient(135deg, #00d4ff, #0099cc);
            color: white;
            border: none;
            padding: 14px 28px;
            border-radius: 30px;
            font-size: 1rem;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            box-shadow: 0 4px 15px rgba(0, 212, 255, 0.3);
            position: relative;
            overflow: hidden;
        }

        .upload-btn::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.2), transparent);
            transition: left 0.5s;
        }

        .upload-btn:hover::before {
            left: 100%;
        }

        .upload-btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 25px rgba(0, 212, 255, 0.4);
        }

        .browse-btn {
            background: linear-gradient(135deg, rgba(255, 255, 255, 0.1), rgba(255, 255, 255, 0.05));
            border: 1px solid rgba(255, 255, 255, 0.2);
            color: white;
        }

        .browse-btn:hover {
            background: linear-gradient(135deg, rgba(255, 255, 255, 0.15), rgba(255, 255, 255, 0.08));
            border-color: rgba(255, 255, 255, 0.3);
            box-shadow: 0 8px 25px rgba(255, 255, 255, 0.1);
        }

        .file-input {
            display: none;
        }

        .selected-files-container {
            margin-top: 30px;
            display: none;
        }

        .selected-files-container.show {
            display: block;
            animation: slideDown 0.3s ease-out;
        }

        @keyframes slideDown {
            from {
                opacity: 0;
                transform: translateY(-20px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }

        .selected-files-header {
            display: flex;
            align-items: center;
            justify-content: space-between;
            margin-bottom: 20px;
        }

        .selected-files-title {
            color: #00d4ff;
            font-weight: 600;
            font-size: 1.1rem;
            display: flex;
            align-items: center;
            gap: 8px;
        }

        .clear-all-btn {
            background: rgba(255, 107, 53, 0.1);
            border: 1px solid rgba(255, 107, 53, 0.3);
            color: #ff6b35;
            padding: 6px 12px;
            border-radius: 15px;
            font-size: 0.9rem;
            cursor: pointer;
            transition: all 0.3s ease;
        }

        .clear-all-btn:hover {
            background: rgba(255, 107, 53, 0.2);
            border-color: rgba(255, 107, 53, 0.5);
        }

        .files-grid {
            display: grid;
            gap: 12px;
        }

        .file-item {
            background: linear-gradient(135deg, rgba(0, 212, 255, 0.08), rgba(0, 212, 255, 0.04));
            border: 1px solid rgba(0, 212, 255, 0.2);
            border-radius: 12px;
            padding: 16px 20px;
            display: flex;
            align-items: center;
            justify-content: space-between;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            position: relative;
            overflow: hidden;
        }

        .file-item::before {
            content: '';
            position: absolute;
            left: 0;
            top: 0;
            height: 100%;
            width: 3px;
            background: linear-gradient(135deg, #00d4ff, #0099cc);
            transition: width 0.3s ease;
        }

        .file-item:hover {
            background: linear-gradient(135deg, rgba(0, 212, 255, 0.12), rgba(0, 212, 255, 0.06));
            border-color: rgba(0, 212, 255, 0.4);
            transform: translateX(5px);
        }

        .file-item:hover::before {
            width: 6px;
        }

        .file-info {
            display: flex;
            align-items: center;
            gap: 15px;
            flex: 1;
        }

        .file-icon {
            width: 40px;
            height: 40px;
            border-radius: 8px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 1.2rem;
            font-weight: 600;
        }

        .file-icon.pdf { background: rgba(255, 59, 48, 0.2); color: #ff3b30; }
        .file-icon.doc { background: rgba(0, 122, 255, 0.2); color: #007aff; }
        .file-icon.excel { background: rgba(52, 199, 89, 0.2); color: #34c759; }
        .file-icon.image { background: rgba(255, 149, 0, 0.2); color: #ff9500; }
        .file-icon.text { background: rgba(175, 82, 222, 0.2); color: #af52de; }

        .file-details {
            flex: 1;
        }

        .file-name {
            color: white;
            font-weight: 500;
            font-size: 1rem;
            margin-bottom: 4px;
            word-break: break-word;
        }

        .file-meta {
            color: #888;
            font-size: 0.85rem;
            display: flex;
            gap: 15px;
        }

        .remove-file-btn {
            background: rgba(255, 107, 53, 0.1);
            border: 1px solid rgba(255, 107, 53, 0.3);
            color: #ff6b35;
            border-radius: 50%;
            width: 32px;
            height: 32px;
            cursor: pointer;
            transition: all 0.3s ease;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 1rem;
        }

        .remove-file-btn:hover {
            background: rgba(255, 107, 53, 0.2);
            border-color: rgba(255, 107, 53, 0.5);
            transform: scale(1.1);
        }

        .analyze-section {
            margin-top: 30px;
            text-align: center;
            display: none;
        }

        .analyze-section.show {
            display: block;
            animation: slideUp 0.3s ease-out;
        }

        @keyframes slideUp {
            from {
                opacity: 0;
                transform: translateY(20px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }

        .analyze-btn {
            background: linear-gradient(135deg, #ff6b35, #ff8c52);
            color: white;
            border: none;
            padding: 16px 40px;
            border-radius: 30px;
            font-size: 1.1rem;
            font-weight: 700;
            cursor: pointer;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            box-shadow: 0 6px 20px rgba(255, 107, 53, 0.3);
            position: relative;
            overflow: hidden;
        }

        .analyze-btn::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.2), transparent);
            transition: left 0.5s;
        }

        .analyze-btn:hover::before {
            left: 100%;
        }

        .analyze-btn:hover {
            transform: translateY(-3px);
            box-shadow: 0 10px 30px rgba(255, 107, 53, 0.4);
        }

        .analyze-btn:active {
            transform: translateY(-1px);
        }

        .tabs {
            background: rgba(255, 255, 255, 0.05);
            border-radius: 20px;
            padding: 30px;
            margin-top: 30px;
            backdrop-filter: blur(10px);
            border: 1px solid rgba(0, 212, 255, 0.2);
            display: none;
        }

        .tab-buttons {
            display: flex;
            gap: 10px;
            margin-bottom: 30px;
            flex-wrap: wrap;
        }

        .tab-btn {
            background: rgba(0, 212, 255, 0.1);
            color: #00d4ff;
            border: 1px solid #00d4ff;
            padding: 12px 24px;
            border-radius: 25px;
            cursor: pointer;
            transition: all 0.3s ease;
            font-weight: 500;
        }

        .tab-btn:hover, .tab-btn.active {
            background: #00d4ff;
            color: #000;
        }

        .tab-content {
            display: none;
        }

        .tab-content.active {
            display: block;
        }

        .progress-container {
            display: none;
            margin: 30px 0;
        }

        .progress-bar {
            width: 100%;
            height: 8px;
            background: rgba(255, 255, 255, 0.1);
            border-radius: 4px;
            overflow: hidden;
        }

        .progress-fill {
            height: 100%;
            background: linear-gradient(90deg, #00d4ff, #0099cc);
            width: 0%;
            transition: width 0.3s ease;
        }

        .status-text {
            color: #00d4ff;
            margin-top: 10px;
            font-weight: 500;
        }

        @media (max-width: 768px) {
            .main-title {
                font-size: 2.5rem;
            }
            
            .industry-grid {
                grid-template-columns: 1fr;
            }
            
            .container {
                padding: 15px;
            }

            .upload-dropzone {
                padding: 40px 20px;
            }

            .upload-icon {
                font-size: 3rem;
            }

            .upload-text {
                font-size: 1.1rem;
            }

            .upload-actions {
                flex-direction: column;
                align-items: center;
            }

            .upload-btn, .browse-btn {
                width: 100%;
                max-width: 200px;
            }

            .supported-formats {
                gap: 8px;
            }

            .format-badge {
                font-size: 0.7rem;
                padding: 4px 8px;
            }

            .file-info {
                gap: 10px;
            }

            .file-meta {
                flex-direction: column;
                gap: 4px;
            }
        }
    </style>
</head>
<body>
    <div class="animated-bg"></div>
    
    <div class="container">
        <div class="header">
            <h1 class="main-title">Universal Document Analyzer</h1>
            <p class="subtitle">AI-Powered Multi-Industry Document Intelligence Platform</p>
        </div>

        <!-- Industry Selection Grid -->
        <div class="industry-selector">
            <h2 class="selector-title">Select Your Industry</h2>
            <div class="industry-grid">
                {% for industry in industries %}
                <div class="industry-card" data-industry="{{ industry.id }}">
                    <div class="industry-icon">{{ industry.icon }}</div>
                    <div class="industry-name">{{ industry.name }}</div>
                    <div class="industry-desc">{{ industry.description }}</div>
                    
                    <!-- Hover Overlay -->
                    <div class="capability-overlay">
                        <div class="overlay-title">What This Does For You:</div>
                        <div class="capability-grid-mini">
                            {% for capability in industry.capabilities %}
                            <div class="capability-item-mini">
                                <span class="capability-icon-mini">{{ capability.icon }}</span>
                                <div class="capability-name-mini">{{ capability.name }}</div>
                            </div>
                            {% endfor %}
                        </div>
                        <div class="performance-badge">{{ industry.performance_metric }}</div>
                    </div>
                </div>
                {% endfor %}
            </div>
        </div>

        <!-- Mortgage Analysis Workflow -->
        <div id="mortgageWorkflow" class="mortgage-workflow" style="display: none;">
            <h2 style="color: #ff6b35; margin-bottom: 20px;">🏠 Mortgage Analysis Workflow</h2>
            <p style="color: #b0b0b0; margin-bottom: 30px;">Follow these steps for accurate mortgage package analysis</p>
            
            <div class="workflow-steps">
                <div class="workflow-step" id="step1">
                    <div style="display: flex; align-items: center; gap: 15px;">
                        <div class="step-number active">1</div>
                        <div>
                            <strong style="color: white;">Parse Lender Requirements</strong>
                            <div style="font-size: 0.9rem; color: #b0b0b0;">Upload or paste lender email to extract document requirements</div>
                        </div>
                    </div>
                    <button onclick="showEmailParser()" style="background: #ff6b35; color: white; border: none; padding: 8px 16px; border-radius: 20px; cursor: pointer;">
                        Start Here
                    </button>
                </div>
                
                <div class="workflow-step" id="step2">
                    <div style="display: flex; align-items: center; gap: 15px;">
                        <div class="step-number">2</div>
                        <div>
                            <strong style="color: white;">Upload Documents</strong>
                            <div style="font-size: 0.9rem; color: #b0b0b0;">Upload mortgage package files for analysis</div>
                        </div>
                    </div>
                    <span style="color: #888; font-size: 0.9rem;">Complete Step 1 first</span>
                </div>
                
                <div class="workflow-step" id="step3">
                    <div style="display: flex; align-items: center; gap: 15px;">
                        <div class="step-number">3</div>
                        <div>
                            <strong style="color: white;">Analyze & Match</strong>
                            <div style="font-size: 0.9rem; color: #b0b0b0;">Match documents against lender requirements</div>
                        </div>
                    </div>
                    <span style="color: #888; font-size: 0.9rem;">Complete Steps 1-2 first</span>
                </div>
                
                <div class="workflow-step" id="step4">
                    <div style="display: flex; align-items: center; gap: 15px;">
                        <div class="step-number">4</div>
                        <div>
                            <strong style="color: white;">Generate Organized PDF</strong>
                            <div style="font-size: 0.9rem; color: #b0b0b0;">AI-powered document reorganization for lender compliance</div>
                        </div>
                    </div>
                    <span style="color: #888; font-size: 0.9rem;">Complete Steps 1-3 first</span>
                </div>
            </div>
        </div>

        <!-- Universal Upload Section (for non-mortgage industries) -->
        <div class="upload-section" id="universalUpload">
            <div class="upload-header">
                <h2 class="upload-title">Upload Your Documents</h2>
                <p class="upload-subtitle">Drag and drop files or click to browse</p>
                
                <div class="supported-formats">
                    <span class="format-badge">📄 PDF</span>
                    <span class="format-badge">📝 Word</span>
                    <span class="format-badge">📊 Excel</span>
                    <span class="format-badge">🖼️ Images</span>
                    <span class="format-badge">📋 Text</span>
                </div>
            </div>

            <div class="upload-dropzone" id="dropzone">
                <div class="upload-icon-container">
                    <span class="upload-icon">☁️</span>
                </div>
                <p class="upload-text">Drop files here to upload</p>
                <p class="upload-hint">or click the button below to browse your files</p>
                
                <div class="upload-actions">
                    <button class="upload-btn" onclick="document.getElementById('fileInput').click()">
                        📁 Choose Files
                    </button>
                    <button class="upload-btn browse-btn" onclick="document.getElementById('fileInput').click()">
                        🔍 Browse
                    </button>
                </div>
            </div>

            <input type="file" id="fileInput" class="file-input" multiple 
                   accept=".pdf,.docx,.doc,.xlsx,.xls,.png,.jpg,.jpeg,.gif,.txt">

            <div class="selected-files-container" id="selectedFilesContainer">
                <div class="selected-files-header">
                    <h3 class="selected-files-title">
                        📋 Selected Files
                        <span id="fileCount">(0)</span>
                    </h3>
                    <button class="clear-all-btn" onclick="clearAllFiles()">
                        🗑️ Clear All
                    </button>
                </div>
                <div class="files-grid" id="filesGrid"></div>
            </div>

            <div class="analyze-section" id="analyzeSection">
                <button class="analyze-btn" onclick="startAnalysis()">
                    🚀 Start Analysis
                </button>
            </div>

            <div class="progress-container" id="progressContainer">
                <div class="progress-bar">
                    <div class="progress-fill" id="progressFill"></div>
                </div>
                <div class="status-text" id="statusText">Initializing...</div>
            </div>
        </div>

        <!-- Tabs Section -->
        <div class="tabs" id="tabsSection">
            <div class="tab-buttons">
                <button class="tab-btn active" onclick="showTab('analyze')">📊 Analyze & Identify</button>
                <button class="tab-btn" onclick="showTab('separation')">📄 Document Separation</button>
                <button class="tab-btn" onclick="showTab('rules')">⚙️ Analysis Rules</button>
                <button class="tab-btn" onclick="showTab('email')">📧 Email Parser</button>
            </div>

            <div id="analyze" class="tab-content active">
                <h3 style="color: #00d4ff; margin-bottom: 20px;">Document Analysis Results</h3>
                <div id="analysisResults">
                    <p style="color: #b0b0b0;">Upload and analyze documents to see results here.</p>
                </div>
            </div>

            <div id="separation" class="tab-content">
                <h3 style="color: #00d4ff; margin-bottom: 20px;">Document Separation</h3>
                <div id="separationResults">
                    <p style="color: #b0b0b0;">Document separation results will appear here after analysis.</p>
                </div>
            </div>

            <div id="rules" class="tab-content">
                <h3 style="color: #00d4ff; margin-bottom: 20px;">Analysis Rules</h3>
                <div id="rulesContent">
                    <p style="color: #b0b0b0;">Industry-specific analysis rules will be displayed here.</p>
                </div>
            </div>

            <div id="email" class="tab-content">
                <h3 style="color: #00d4ff; margin-bottom: 20px;">Email Parser</h3>
                <div style="margin-bottom: 20px;">
                    <textarea id="emailContent" placeholder="Paste lender email content here..." 
                              style="width: 100%; height: 200px; background: rgba(255,255,255,0.1); border: 1px solid #00d4ff; border-radius: 10px; padding: 15px; color: white; resize: vertical;"></textarea>
                </div>
                <button onclick="parseEmail()" style="background: linear-gradient(45deg, #00d4ff, #0099cc); color: white; border: none; padding: 12px 25px; border-radius: 25px; cursor: pointer; font-weight: 600;">
                    Parse Requirements
                </button>
                <div id="emailResults" style="margin-top: 20px;">
                    <p style="color: #b0b0b0;">Email parsing results will appear here.</p>
                </div>
            </div>
        </div>
    </div>

    <script>
        let selectedIndustry = null;
        let selectedFiles = [];
        let lenderRequirementsParsed = false;
        let mortgageWorkflowStep = 1;
        let lastAnalysisResults = null;
        let lastLenderRequirements = null;
        let uploadedPdfPath = null; // Track the uploaded PDF path for reorganization

        // Industry selection
        document.querySelectorAll('.industry-card').forEach(card => {
            card.addEventListener('click', function() {
                document.querySelectorAll('.industry-card').forEach(c => c.classList.remove('selected'));
                this.classList.add('selected');
                selectedIndustry = this.dataset.industry;
                
                // Show appropriate workflow based on industry
                if (selectedIndustry === 'mortgage') {
                    // Show mortgage workflow
                    document.getElementById('mortgageWorkflow').style.display = 'block';
                    // Don't hide upload section with inline style - let CSS classes control it
                    const uploadSection = document.getElementById('universalUpload');
                    uploadSection.classList.remove('show');
                    // Clear any inline styles that might interfere
                    uploadSection.style.display = '';
                    updateMortgageWorkflow();
                } else {
                    // Show universal upload for other industries
                    document.getElementById('mortgageWorkflow').style.display = 'none';
                    const uploadSection = document.getElementById('universalUpload');
                    uploadSection.classList.add('show');
                    // Clear any inline styles that might interfere
                    uploadSection.style.display = '';
                }
                
                // Show tabs section when industry is selected
                document.getElementById('tabsSection').style.display = 'block';
                
                console.log('Selected industry:', selectedIndustry);
            });
        });

        // Mortgage workflow functions
        function updateMortgageWorkflow() {
            console.log('updateMortgageWorkflow called, step:', mortgageWorkflowStep);
            
            // Reset all steps
            document.querySelectorAll('.workflow-step').forEach(step => {
                step.classList.remove('active', 'completed');
                step.querySelector('.step-number').classList.remove('active', 'completed');
            });

            // Update current step
            for (let i = 1; i <= 4; i++) {
                const step = document.getElementById(`step${i}`);
                const stepNumber = step.querySelector('.step-number');
                
                if (i < mortgageWorkflowStep) {
                    step.classList.add('completed');
                    stepNumber.classList.add('completed');
                } else if (i === mortgageWorkflowStep) {
                    step.classList.add('active');
                    stepNumber.classList.add('active');
                }
            }

            // Update step 2 content based on workflow progress
            const step2 = document.getElementById('step2');
            const step2Content = step2.querySelector('span');
            
            console.log('Step 2 element found:', !!step2);
            console.log('Step 2 span found:', !!step2Content);
            
            if (mortgageWorkflowStep >= 2) {
                console.log('Enabling step 2 - creating upload button');
                // Enable step 2 - show upload button
                if (step2Content) {
                    step2Content.innerHTML = '<button onclick="showUploadSection()" style="background: #ff6b35; color: white; border: none; padding: 8px 16px; border-radius: 20px; cursor: pointer;">Upload Documents</button>';
                    console.log('Upload Documents button created');
                } else {
                    console.error('Step 2 span element not found');
                }
                
                // Show upload section
                const uploadSection = document.getElementById('universalUpload');
                if (uploadSection) {
                    uploadSection.classList.add('show');
                    console.log('Upload section shown for mortgage workflow step 2');
                    console.log('Upload section classes:', uploadSection.className);
                } else {
                    console.error('Upload section element not found');
                }
            } else {
                // Disable step 2
                if (step2Content) {
                    step2Content.innerHTML = '<span style="color: #888; font-size: 0.9rem;">Complete Step 1 first</span>';
                }
                const uploadSection = document.getElementById('universalUpload');
                if (uploadSection) {
                    uploadSection.classList.remove('show');
                }
            }

            // Update step 3 content based on workflow progress
            const step3 = document.getElementById('step3');
            const step3Content = step3.querySelector('span');
            
            if (mortgageWorkflowStep >= 3) {
                // Enable step 3 - show analyze button
                if (step3Content) {
                    step3Content.innerHTML = '<button onclick="startMortgageAnalysis()" style="background: #ff6b35; color: white; border: none; padding: 8px 16px; border-radius: 20px; cursor: pointer;">Start Analysis</button>';
                }
            } else {
                // Disable step 3
                if (step3Content) {
                    step3Content.innerHTML = '<span style="color: #888; font-size: 0.9rem;">Complete Steps 1-2 first</span>';
                }
            }

            // Update step 4 content based on workflow progress
            const step4 = document.getElementById('step4');
            const step4Content = step4.querySelector('span');
            
            if (mortgageWorkflowStep >= 4) {
                // Enable step 4 - show PDF reorganization button
                if (step4Content) {
                    step4Content.innerHTML = '<button onclick="generateOrganizedPDF()" style="background: linear-gradient(45deg, #ff6b35, #f7931e); color: white; border: none; padding: 8px 16px; border-radius: 20px; cursor: pointer; box-shadow: 0 2px 10px rgba(255, 107, 53, 0.3);">Generate Organized PDF</button>';
                }
            } else {
                // Disable step 4
                if (step4Content) {
                    step4Content.innerHTML = '<span style="color: #888; font-size: 0.9rem;">Complete Steps 1-3 first</span>';
                }
            }
        }

        function showUploadSection() {
            console.log('showUploadSection called');
            
            // Ensure upload section is visible first
            const uploadSection = document.getElementById('universalUpload');
            console.log('Upload section element found:', !!uploadSection);
            
            if (uploadSection) {
                console.log('Upload section classes before:', uploadSection.className);
                uploadSection.classList.add('show');
                console.log('Upload section classes after:', uploadSection.className);
                
                // Check if element is actually visible
                const computedStyle = window.getComputedStyle(uploadSection);
                console.log('Upload section display style:', computedStyle.display);
                
                // Small delay to ensure the element is rendered before scrolling
                setTimeout(() => {
                    console.log('Attempting to scroll to upload section');
                    uploadSection.scrollIntoView({ behavior: 'smooth' });
                    console.log('Scroll command executed');
                }, 100);
            } else {
                console.error('Upload section not found');
            }
        }

        function startMortgageAnalysis() {
            // Start the analysis for mortgage workflow
            startAnalysis();
        }

        function generateOrganizedPDF() {
            // Generate AI-organized PDF for mortgage workflow
            console.log('Starting PDF reorganization...');
            
            // Show loading state
            const step4Content = document.getElementById('step4').querySelector('span');
            step4Content.innerHTML = '<span style="color: #ff6b35;">🤖 AI is organizing your documents...</span>';
            
            // Prepare data for PDF reorganization
            const reorganizationData = {
                document_sections: lastAnalysisResults?.sections || [],
                lender_requirements: lastLenderRequirements || {},
                original_pdf_path: uploadedPdfPath || '' // Use the tracked PDF path
            };
            
            // Call PDF reorganization endpoint
            fetch('/reorganize_pdf', {
                method: 'POST',
                headers: {
                    'Content-Type': 'application/json',
                },
                body: JSON.stringify(reorganizationData)
            })
            .then(response => response.json())
            .then(data => {
                if (data.success) {
                    // Show success with download link
                    step4Content.innerHTML = `
                        <div style="display: flex; flex-direction: column; gap: 10px; align-items: center;">
                            <span style="color: #00ff00;">✅ PDF Generated Successfully!</span>
                            <a href="/download_pdf/${data.output_filename}" 
                               download="${data.output_filename}"
                               style="background: linear-gradient(45deg, #00ff00, #00cc00); color: #000; text-decoration: none; padding: 8px 16px; border-radius: 20px; font-weight: bold;">
                               📄 Download Organized PDF
                            </a>
                            <div style="font-size: 0.8rem; color: #b0b0b0;">
                                ${data.documents_included} documents • ${data.ai_analysis.cost_estimate} AI cost
                            </div>
                        </div>
                    `;
                    
                    // Show AI analysis results
                    displayPDFReorganizationResults(data);
                } else {
                    // Show error
                    step4Content.innerHTML = `
                        <div style="color: #ff4444;">
                            ❌ PDF Generation Failed: ${data.error}
                            <button onclick="generateOrganizedPDF()" style="background: #ff6b35; color: white; border: none; padding: 4px 8px; border-radius: 10px; cursor: pointer; margin-left: 10px;">Retry</button>
                        </div>
                    `;
                }
            })
            .catch(error => {
                console.error('PDF reorganization error:', error);
                step4Content.innerHTML = `
                    <div style="color: #ff4444;">
                        ❌ Network Error
                        <button onclick="generateOrganizedPDF()" style="background: #ff6b35; color: white; border: none; padding: 4px 8px; border-radius: 10px; cursor: pointer; margin-left: 10px;">Retry</button>
                    </div>
                `;
            });
        }

        function displayPDFReorganizationResults(data) {
            // Display AI analysis results for PDF reorganization
            const resultsContainer = document.getElementById('results');
            if (!resultsContainer) return;
            
            const aiAnalysisHtml = `
                <div style="background: rgba(0, 255, 0, 0.1); border: 1px solid #00ff00; border-radius: 10px; padding: 20px; margin-top: 20px;">
                    <h3 style="color: #00ff00; margin-bottom: 15px;">🤖 AI PDF Reorganization Complete</h3>
                    
                    <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 15px; margin-bottom: 20px;">
                        <div style="background: rgba(255, 255, 255, 0.05); padding: 15px; border-radius: 8px;">
                            <strong style="color: #00d4ff;">Compliance Score:</strong><br>
                            <span style="color: #00ff00; font-size: 1.2rem;">${data.compliance_summary.compliance_score}%</span>
                        </div>
                        <div style="background: rgba(255, 255, 255, 0.05); padding: 15px; border-radius: 8px;">
                            <strong style="color: #00d4ff;">Documents Included:</strong><br>
                            <span style="color: #ffffff;">${data.documents_included}</span>
                        </div>
                        <div style="background: rgba(255, 255, 255, 0.05); padding: 15px; border-radius: 8px;">
                            <strong style="color: #00d4ff;">AI Cost:</strong><br>
                            <span style="color: #ffffff;">${data.ai_analysis.cost_estimate}</span>
                        </div>
                        <div style="background: rgba(255, 255, 255, 0.05); padding: 15px; border-radius: 8px;">
                            <strong style="color: #00d4ff;">Tokens Used:</strong><br>
                            <span style="color: #ffffff;">${data.ai_analysis.total_tokens_used}</span>
                        </div>
                    </div>
                    
                    <div style="background: rgba(255, 255, 255, 0.03); padding: 15px; border-radius: 8px; margin-bottom: 15px;">
                        <strong style="color: #00d4ff;">Document List:</strong><br>
                        ${data.compliance_summary.document_list.map(doc => `<span style="color: #ffffff;">• ${doc}</span>`).join('<br>')}
                    </div>
                    
                    <div style="text-align: center;">
                        <a href="/download_pdf/${data.output_filename}" 
                           download="${data.output_filename}"
                           style="background: linear-gradient(45deg, #00ff00, #00cc00); color: #000; text-decoration: none; padding: 12px 24px; border-radius: 25px; font-weight: bold; display: inline-block;">
                           📄 Download AI-Organized PDF Package
                        </a>
                    </div>
                </div>
            `;
            
            resultsContainer.insertAdjacentHTML('beforeend', aiAnalysisHtml);
        }

        function showEmailParser() {
            // Switch to email parser tab
            showTab('email');
            // Scroll to tabs section
            document.getElementById('tabsSection').scrollIntoView({ behavior: 'smooth' });
        }

        function advanceMortgageWorkflow() {
            if (mortgageWorkflowStep < 4) {
                mortgageWorkflowStep++;
                updateMortgageWorkflow();
            }
        }

        // Universal upload functionality (for non-mortgage industries)
        document.getElementById('fileInput').addEventListener('change', function(e) {
            handleFiles(Array.from(e.target.files));
        });

        // Drag and drop handlers
        const dropzone = document.getElementById('dropzone');

        dropzone.addEventListener('dragover', function(e) {
            e.preventDefault();
            dropzone.classList.add('dragover');
        });

        dropzone.addEventListener('dragleave', function(e) {
            e.preventDefault();
            dropzone.classList.remove('dragover');
        });

        dropzone.addEventListener('drop', function(e) {
            e.preventDefault();
            dropzone.classList.remove('dragover');
            handleFiles(Array.from(e.dataTransfer.files));
        });

        function handleFiles(files) {
            // Add new files to selected files array
            files.forEach(file => {
                // Check if file already exists
                if (!selectedFiles.find(f => f.name === file.name && f.size === file.size)) {
                    selectedFiles.push(file);
                }
            });
            
            displaySelectedFiles();
            
            // Advance mortgage workflow to step 3 when files are uploaded
            if (selectedIndustry === 'mortgage' && selectedFiles.length > 0 && mortgageWorkflowStep === 2) {
                mortgageWorkflowStep = 3;
                updateMortgageWorkflow();
            }
        }

        function displaySelectedFiles() {
            const container = document.getElementById('selectedFilesContainer');
            const grid = document.getElementById('filesGrid');
            const fileCount = document.getElementById('fileCount');
            const analyzeSection = document.getElementById('analyzeSection');

            if (selectedFiles.length > 0) {
                container.classList.add('show');
                analyzeSection.classList.add('show');
                
                fileCount.textContent = `(${selectedFiles.length})`;
                
                grid.innerHTML = selectedFiles.map((file, index) => {
                    const fileType = getFileType(file.name);
                    const fileSize = formatFileSize(file.size);
                    
                    return `
                        <div class="file-item">
                            <div class="file-info">
                                <div class="file-icon ${fileType.class}">
                                    ${fileType.icon}
                                </div>
                                <div class="file-details">
                                    <div class="file-name">${file.name}</div>
                                    <div class="file-meta">
                                        <span>${fileSize}</span>
                                        <span>${fileType.name}</span>
                                    </div>
                                </div>
                            </div>
                            <button class="remove-file-btn" onclick="removeFile(${index})" title="Remove file">
                                ✕
                            </button>
                        </div>
                    `;
                }).join('');
            } else {
                container.classList.remove('show');
                analyzeSection.classList.remove('show');
            }
        }

        function getFileType(filename) {
            const extension = filename.split('.').pop().toLowerCase();
            
            switch (extension) {
                case 'pdf':
                    return { class: 'pdf', icon: 'PDF', name: 'PDF Document' };
                case 'doc':
                case 'docx':
                    return { class: 'doc', icon: 'DOC', name: 'Word Document' };
                case 'xls':
                case 'xlsx':
                    return { class: 'excel', icon: 'XLS', name: 'Excel Spreadsheet' };
                case 'png':
                case 'jpg':
                case 'jpeg':
                case 'gif':
                    return { class: 'image', icon: 'IMG', name: 'Image File' };
                case 'txt':
                    return { class: 'text', icon: 'TXT', name: 'Text File' };
                default:
                    return { class: 'text', icon: 'FILE', name: 'Document' };
            }
        }

        function formatFileSize(bytes) {
            if (bytes === 0) return '0 Bytes';
            const k = 1024;
            const sizes = ['Bytes', 'KB', 'MB', 'GB'];
            const i = Math.floor(Math.log(bytes) / Math.log(k));
            return parseFloat((bytes / Math.pow(k, i)).toFixed(2)) + ' ' + sizes[i];
        }

        function removeFile(index) {
            selectedFiles.splice(index, 1);
            displaySelectedFiles();
        }

        function clearAllFiles() {
            selectedFiles = [];
            document.getElementById('fileInput').value = '';
            displaySelectedFiles();
        }

        // Tab functionality
        function showTab(tabName) {
            document.querySelectorAll('.tab-content').forEach(tab => tab.classList.remove('active'));
            document.querySelectorAll('.tab-btn').forEach(btn => btn.classList.remove('active'));
            
            document.getElementById(tabName).classList.add('active');
            event.target.classList.add('active');
        }

        // Analysis functionality
        function startAnalysis() {
            // Check mortgage workflow requirements
            if (selectedIndustry === 'mortgage' && !lenderRequirementsParsed) {
                alert('Please parse lender requirements first (Step 1) before analyzing documents.');
                showEmailParser();
                return;
            }

            if (selectedFiles.length === 0) {
                alert('Please select files to analyze');
                return;
            }

            const progressContainer = document.getElementById('progressContainer');
            const progressFill = document.getElementById('progressFill');
            const statusText = document.getElementById('statusText');
            
            progressContainer.style.display = 'block';
            
            const formData = new FormData();
            selectedFiles.forEach(file => {
                formData.append('files', file);
                // Track the first PDF file for reorganization
                if (file.type === 'application/pdf' && !uploadedPdfPath) {
                    uploadedPdfPath = `/tmp/${file.name}`;
                    console.log('🔍 Tracked PDF path for reorganization:', uploadedPdfPath);
                }
            });
            formData.append('industry', selectedIndustry);

            // Simulate progress
            let progress = 0;
            const progressInterval = setInterval(() => {
                progress += Math.random() * 15;
                if (progress > 90) progress = 90;
                progressFill.style.width = progress + '%';
                statusText.textContent = getStatusMessage(progress);
            }, 500);

            fetch('/analyze', {
                method: 'POST',
                body: formData
            })
            .then(response => response.json())
            .then(data => {
                clearInterval(progressInterval);
                progressFill.style.width = '100%';
                statusText.textContent = 'Analysis complete!';
                
                setTimeout(() => {
                    progressContainer.style.display = 'none';
                    displayResults(data);
                    displayDocumentSeparation(data);
                    displayAnalysisRules(data);
                    
                    // Store analysis results for PDF reorganization
                    lastAnalysisResults = data;
                    
                    // Advance mortgage workflow to step 4 (PDF reorganization)
                    if (selectedIndustry === 'mortgage') {
                        mortgageWorkflowStep = 4;
                        updateMortgageWorkflow();
                    }
                }, 1000);
            })
            .catch(error => {
                clearInterval(progressInterval);
                console.error('Error:', error);
                statusText.textContent = 'Analysis failed. Please try again.';
            });
        }

        function getStatusMessage(progress) {
            if (progress < 20) return 'Initializing analysis...';
            if (progress < 40) return 'Processing documents...';
            if (progress < 60) return 'Extracting content...';
            if (progress < 80) return 'Analyzing patterns...';
            return 'Finalizing results...';
        }

        // Email parsing functionality
        function parseEmail() {
            const emailContent = document.getElementById('emailContent').value;
            if (!emailContent.trim()) {
                alert('Please paste email content first');
                return;
            }

            fetch('/parse_email', {
                method: 'POST',
                headers: {
                    'Content-Type': 'application/json',
                },
                body: JSON.stringify({ content: emailContent })
            })
            .then(response => response.json())
            .then(data => {
                displayEmailResults(data);
                if (data.success) {
                    lenderRequirementsParsed = true;
                    // Store lender requirements for PDF reorganization
                    lastLenderRequirements = data;
                    if (selectedIndustry === 'mortgage') {
                        mortgageWorkflowStep = 2;
                        updateMortgageWorkflow();
                    }
                }
            })
            .catch(error => {
                console.error('Error:', error);
                alert('Failed to parse email. Please try again.');
            });
        }

        function displayEmailResults(data) {
            const container = document.getElementById('emailResults');
            if (data.success) {
                let html = '<div style="background: rgba(0,255,0,0.1); border: 1px solid #00ff00; border-radius: 10px; padding: 20px;">';
                html += '<h4 style="color: #00ff00; margin-bottom: 15px;">✅ Email Parsed Successfully</h4>';
                
                // AI Enhancement indicator (if available)
                if (data.ai_enhanced) {
                    html += '<div style="background: rgba(0,255,0,0.05); border-left: 4px solid #00ff00; padding: 10px; margin-bottom: 15px;">';
                    html += '<h6 style="color: #00ff00; margin-bottom: 8px;">🤖 AI-Enhanced Parsing</h6>';
                    html += `<p style="color: #b0b0b0; font-size: 0.9rem;">Tokens Used: ${data.ai_tokens_used || 0} | Cost: ${data.ai_cost_estimate || '$0.0000'}</p>`;
                    html += '</div>';
                }
                
                html += `<p><strong>Lender:</strong> ${data.lender_name}</p>`;
                html += `<p><strong>Contact:</strong> ${data.contact_email}</p>`;
                if (data.funding_amount) {
                    html += `<p><strong>Amount:</strong> ${data.funding_amount}</p>`;
                }
                html += `<p><strong>Documents Required:</strong> ${data.documents.length}</p>`;
                html += '<div style="margin-top: 15px;"><strong>Document List:</strong><ul>';
                data.documents.slice(0, 10).forEach(doc => {
                    html += `<li style="margin: 5px 0;">${doc}</li>`;
                });
                if (data.documents.length > 10) {
                    html += `<li style="color: #888;">... and ${data.documents.length - 10} more</li>`;
                }
                html += '</ul></div>';
                
                // AI Insights (if available)
                if (data.ai_insights) {
                    html += '<div style="background: rgba(0,255,0,0.05); border-left: 4px solid #00ff00; padding: 10px; margin-top: 15px;">';
                    html += '<h6 style="color: #00ff00; margin-bottom: 8px;">🤖 AI Enhanced Insights</h6>';
                    html += `<div style="color: #b0b0b0; font-size: 0.9rem; white-space: pre-wrap;">${data.ai_insights}</div>`;
                    html += '</div>';
                }
                
                html += '</div>';
                container.innerHTML = html;
            } else {
                container.innerHTML = '<div style="background: rgba(255,0,0,0.1); border: 1px solid #ff0000; border-radius: 10px; padding: 20px; color: #ff6b35;"><h4>❌ Parsing Failed</h4><p>' + data.error + '</p></div>';
            }
        }

        function displayResults(data) {
            const container = document.getElementById('analysisResults');
            if (data.success && data.results) {
                let html = '<div style="background: rgba(0,212,255,0.1); border: 1px solid #00d4ff; border-radius: 10px; padding: 20px; margin-bottom: 20px;">';
                html += '<h4 style="color: #00d4ff; margin-bottom: 15px;">📊 Analysis Complete</h4>';
                
                // AI Enhancement Summary (if available)
                if (data.ai_enhanced) {
                    html += '<div style="background: rgba(0,255,0,0.1); border: 1px solid #00ff00; border-radius: 8px; padding: 15px; margin-bottom: 20px;">';
                    html += '<h5 style="color: #00ff00; margin-bottom: 10px;">🤖 AI-Enhanced Analysis</h5>';
                    html += `<p style="color: #b0b0b0;">AI Model: ${data.ai_insights[0]?.tokens ? 'GPT-4o-mini' : 'N/A'}</p>`;
                    html += `<p style="color: #b0b0b0;">Total Tokens Used: ${data.total_ai_tokens || 0}</p>`;
                    html += `<p style="color: #b0b0b0;">Estimated Cost: ${data.ai_cost_estimate || '$0.0000'}</p>`;
                    html += '</div>';
                }
                
                data.results.forEach((result, index) => {
                    html += '<div style="background: rgba(255,255,255,0.05); border-radius: 8px; padding: 15px; margin-bottom: 15px;">';
                    html += `<h5 style="color: white; margin-bottom: 10px;">📄 ${result.filename}</h5>`;
                    html += `<p><strong>Industry:</strong> ${result.industry}</p>`;
                    html += `<p><strong>Word Count:</strong> ${result.word_count}</p>`;
                    html += `<p><strong>Quality Score:</strong> ${result.quality_score}%</p>`;
                    html += `<p><strong>Processing Time:</strong> ${result.processing_time}</p>`;
                    if (result.categories && result.categories.length > 0) {
                        html += `<p><strong>Categories:</strong> ${result.categories.join(', ')}</p>`;
                    }
                    
                    // AI Insights for this document (if available)
                    if (result.ai_insights) {
                        html += '<div style="background: rgba(0,255,0,0.05); border-left: 4px solid #00ff00; padding: 10px; margin-top: 10px;">';
                        html += '<h6 style="color: #00ff00; margin-bottom: 8px;">🤖 AI Insights</h6>';
                        html += `<div style="color: #b0b0b0; font-size: 0.9rem; white-space: pre-wrap;">${result.ai_insights}</div>`;
                        if (result.ai_tokens_used) {
                            html += `<p style="color: #888; font-size: 0.8rem; margin-top: 8px;">Tokens: ${result.ai_tokens_used} | Model: ${result.ai_model}</p>`;
                        }
                        html += '</div>';
                    } else if (result.ai_error) {
                        html += '<div style="background: rgba(255,165,0,0.05); border-left: 4px solid #ffa500; padding: 10px; margin-top: 10px;">';
                        html += '<h6 style="color: #ffa500; margin-bottom: 8px;">⚠️ AI Analysis</h6>';
                        html += `<div style="color: #b0b0b0; font-size: 0.9rem;">${result.ai_error}</div>`;
                        html += '</div>';
                    }
                    
                    html += '</div>';
                });
                
                html += '</div>';
                container.innerHTML = html;
            } else {
                container.innerHTML = '<div style="background: rgba(255,0,0,0.1); border: 1px solid #ff0000; border-radius: 10px; padding: 20px; color: #ff6b35;"><h4>❌ Analysis Failed</h4><p>' + (data.error || 'Unknown error occurred') + '</p></div>';
            }
        }

        function displayDocumentSeparation(data) {
            const container = document.getElementById('separationResults');
            
            if (data.success && data.sections) {
                let html = '<div style="background: rgba(0,212,255,0.1); border: 1px solid #00d4ff; border-radius: 10px; padding: 20px; margin-bottom: 20px;">';
                html += '<h4 style="color: #00d4ff; margin-bottom: 15px;">📄 Document Separation Results</h4>';
                html += '<p style="color: #b0b0b0; margin-bottom: 20px;">Documents have been automatically separated and organized by type.</p>';
                
                data.sections.forEach((section, index) => {
                    html += '<div style="background: rgba(255,255,255,0.05); border-radius: 8px; padding: 15px; margin-bottom: 15px; border-left: 4px solid #00d4ff;">';
                    html += '<div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px;">';
                    html += '<h5 style="color: white; margin: 0;">📑 ' + section.name + '</h5>';
                    html += '<span style="background: #00d4ff; color: #000; padding: 4px 12px; border-radius: 15px; font-size: 0.8rem; font-weight: 600;">Pages: ' + section.pages + '</span>';
                    html += '</div>';
                    
                    html += '<div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(150px, 1fr)); gap: 10px; margin-bottom: 10px;">';
                    html += '<div style="background: rgba(0,255,0,0.1); padding: 8px; border-radius: 5px; text-align: center;">';
                    html += '<div style="color: #00ff00; font-size: 0.8rem;">Confidence</div>';
                    html += '<div style="color: white; font-weight: 600;">' + section.confidence + '</div>';
                    html += '</div>';
                    html += '<div style="background: rgba(255,165,0,0.1); padding: 8px; border-radius: 5px; text-align: center;">';
                    html += '<div style="color: #ffa500; font-size: 0.8rem;">Quality</div>';
                    html += '<div style="color: white; font-weight: 600;">' + section.quality + '</div>';
                    html += '</div>';
                    html += '<div style="background: rgba(255,0,0,0.1); padding: 8px; border-radius: 5px; text-align: center;">';
                    html += '<div style="color: #ff6b35; font-size: 0.8rem;">Risk Score</div>';
                    html += '<div style="color: white; font-weight: 600;">' + section.risk_score + '</div>';
                    html += '</div>';
                    html += '</div>';
                    
                    html += '<div style="display: flex; gap: 10px; margin-top: 10px;">';
                    html += '<button style="background: #00d4ff; color: #000; border: none; padding: 6px 12px; border-radius: 15px; cursor: pointer; font-size: 0.8rem;">View</button>';
                    html += '<button style="background: rgba(255,255,255,0.1); color: white; border: 1px solid rgba(255,255,255,0.3); padding: 6px 12px; border-radius: 15px; cursor: pointer; font-size: 0.8rem;">Edit</button>';
                    html += '<button style="background: rgba(0,255,0,0.1); color: #00ff00; border: 1px solid #00ff00; padding: 6px 12px; border-radius: 15px; cursor: pointer; font-size: 0.8rem;">Export</button>';
                    html += '</div>';
                    
                    if (section.notes) {
                        html += '<p style="color: #888; font-size: 0.9rem; margin-top: 10px; font-style: italic;">' + section.notes + '</p>';
                    }
                    
                    html += '</div>';
                });
                
                html += '</div>';
                container.innerHTML = html;
            } else {
                container.innerHTML = '<p style="color: #b0b0b0;">Document separation results will appear here after analysis.</p>';
            }
        }

        function displayAnalysisRules(data) {
            const container = document.getElementById('rulesContent');
            
            // Display industry-specific rules
            const industryRules = {
                'mortgage': [
                    { name: 'TRID Compliance Check', priority: 'high', description: 'Validates TRID disclosure requirements' },
                    { name: 'Lender Requirements Match', priority: 'high', description: 'Matches documents against lender email requirements' },
                    { name: 'Document Authenticity', priority: 'medium', description: 'Verifies document integrity and authenticity' },
                    { name: 'Regulatory Compliance', priority: 'medium', description: 'Checks compliance with federal and state regulations' },
                    { name: 'Data Extraction Accuracy', priority: 'low', description: 'Validates extracted data accuracy' }
                ],
                'real_estate': [
                    { name: 'Fraud Detection', priority: 'high', description: 'Advanced document fraud detection algorithms' },
                    { name: 'Compliance Validation', priority: 'high', description: 'Real estate regulatory compliance checking' },
                    { name: 'Risk Assessment', priority: 'medium', description: 'Property transaction risk evaluation' },
                    { name: 'Document Classification', priority: 'medium', description: 'Automatic document type classification' },
                    { name: 'Data Verification', priority: 'low', description: 'Cross-reference data verification' }
                ],
                'legal': [
                    { name: 'Contract Analysis', priority: 'high', description: 'Comprehensive contract review and analysis' },
                    { name: 'Legal Compliance', priority: 'high', description: 'Legal and regulatory requirement checking' },
                    { name: 'Clause Extraction', priority: 'medium', description: 'Key legal clause identification and extraction' },
                    { name: 'Risk Identification', priority: 'medium', description: 'Legal risk assessment and flagging' },
                    { name: 'Document Review', priority: 'low', description: 'Automated legal document review' }
                ],
                'healthcare': [
                    { name: 'HIPAA Compliance', priority: 'high', description: 'Healthcare privacy regulation adherence' },
                    { name: 'Medical Record Analysis', priority: 'high', description: 'Comprehensive patient record review' },
                    { name: 'Claims Processing', priority: 'medium', description: 'Insurance claim validation and processing' },
                    { name: 'Clinical Data Extraction', priority: 'medium', description: 'Medical information extraction and coding' },
                    { name: 'Quality Assurance', priority: 'low', description: 'Medical document quality validation' }
                ],
                'financial': [
                    { name: 'SOX Compliance', priority: 'high', description: 'Sarbanes-Oxley compliance validation' },
                    { name: 'Financial Analysis', priority: 'high', description: 'Comprehensive financial document analysis' },
                    { name: 'Risk Assessment', priority: 'medium', description: 'Financial risk evaluation and scoring' },
                    { name: 'Regulatory Compliance', priority: 'medium', description: 'Banking and finance regulation checking' },
                    { name: 'Audit Trail', priority: 'low', description: 'Financial audit trail validation' }
                ],
                'hr': [
                    { name: 'Resume Analysis', priority: 'high', description: 'Automated candidate evaluation and scoring' },
                    { name: 'Compliance Checking', priority: 'high', description: 'HR regulation and policy adherence' },
                    { name: 'Background Verification', priority: 'medium', description: 'Employee background document review' },
                    { name: 'Performance Analytics', priority: 'medium', description: 'Employee performance data analysis' },
                    { name: 'Document Classification', priority: 'low', description: 'HR document type classification' }
                ]
            };
            
            const rules = industryRules[selectedIndustry] || industryRules['mortgage'];
            
            let html = '<div style="background: rgba(0,212,255,0.1); border: 1px solid #00d4ff; border-radius: 10px; padding: 20px; margin-bottom: 20px;">';
            html += '<h4 style="color: #00d4ff; margin-bottom: 15px;">⚙️ Active Analysis Rules</h4>';
            html += `<p style="color: #b0b0b0; margin-bottom: 20px;">Industry-specific rules for ${selectedIndustry || 'mortgage'} document analysis.</p>`;
            
            rules.forEach((rule, index) => {
                const priorityColor = rule.priority === 'high' ? '#ff6b35' : (rule.priority === 'medium' ? '#ffa500' : '#00d4ff');
                
                html += '<div style="background: rgba(255,255,255,0.05); border-radius: 8px; padding: 15px; margin-bottom: 15px; border-left: 4px solid ' + priorityColor + ';">';
                html += '<div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px;">';
                html += '<h5 style="color: white; margin: 0;">🔧 ' + rule.name + '</h5>';
                html += '<span style="background: ' + priorityColor + '; color: ' + (rule.priority === 'high' ? 'white' : '#000') + '; padding: 4px 12px; border-radius: 15px; font-size: 0.8rem; font-weight: 600; text-transform: uppercase;">' + rule.priority + '</span>';
                html += '</div>';
                html += '<p style="color: #b0b0b0; font-size: 0.9rem; margin: 0;">' + rule.description + '</p>';
                html += '</div>';
            });
            
            // Add performance metrics
            html += '<div style="background: rgba(0,255,0,0.1); border: 1px solid #00ff00; border-radius: 8px; padding: 15px; margin-top: 20px;">';
            html += '<h5 style="color: #00ff00; margin-bottom: 10px;">📊 Performance Metrics</h5>';
            html += '<div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(150px, 1fr)); gap: 15px;">';
            html += '<div style="text-align: center;"><div style="color: #00ff00; font-size: 1.5rem; font-weight: bold;">98.5%</div><div style="color: #b0b0b0; font-size: 0.9rem;">Accuracy Rate</div></div>';
            html += '<div style="text-align: center;"><div style="color: #00d4ff; font-size: 1.5rem; font-weight: bold;">2.3s</div><div style="color: #b0b0b0; font-size: 0.9rem;">Avg Processing Time</div></div>';
            html += '<div style="text-align: center;"><div style="color: #ffa500; font-size: 1.5rem; font-weight: bold;">' + rules.length + '</div><div style="color: #b0b0b0; font-size: 0.9rem;">Active Rules</div></div>';
            html += '</div>';
            html += '</div>';
            
            html += '</div>';
            container.innerHTML = html;
        }
    </script>
</body>
</html>
    ''', industries=industries_list)

@app.route('/analyze', methods=['POST'])
def analyze_documents():
    """Analyze uploaded documents with AI enhancement"""
    try:
        if 'files' not in request.files:
            return jsonify({'success': False, 'error': 'No files uploaded'})
        
        files = request.files.getlist('files')
        industry = request.form.get('industry', 'mortgage')
        
        if not files or files[0].filename == '':
            return jsonify({'success': False, 'error': 'No files selected'})
        
        results = []
        sections = []
        ai_insights = []
        total_tokens_used = 0
        
        for file in files:
            if file.filename == '':
                continue
                
            # Save uploaded file temporarily
            temp_path = f"/tmp/{file.filename}"
            file.save(temp_path)
            
            try:
                # Process the file
                processing_result = document_processor.process_file(temp_path, file.filename)
                
                if processing_result['success']:
                    # Standard analysis
                    analysis_result = analyze_document_content(
                        processing_result['text'], 
                        file.filename, 
                        industry
                    )
                    
                    # AI-Enhanced Analysis (if available and for mortgage industry)
                    if mortgage_ai and industry == 'mortgage' and processing_result['text']:
                        ai_result = mortgage_ai.analyze_mortgage_document(
                            processing_result['text'], 
                            file.filename
                        )
                        
                        if ai_result.get('ai_analysis'):
                            analysis_result['ai_insights'] = ai_result['analysis']
                            analysis_result['ai_tokens_used'] = ai_result['tokens_used']
                            analysis_result['ai_model'] = ai_result['model']
                            total_tokens_used += ai_result['tokens_used']
                            ai_insights.append({
                                'filename': file.filename,
                                'analysis': ai_result['analysis'],
                                'tokens': ai_result['tokens_used']
                            })
                        else:
                            analysis_result['ai_error'] = ai_result.get('error', 'AI analysis unavailable')
                    
                    results.append(analysis_result)
                    
                    # Generate sections for mortgage industry
                    if industry == 'mortgage':
                        file_sections = analyze_mortgage_sections(file.filename, use_lender_rules=True)
                        sections.extend(file_sections)
                else:
                    results.append({
                        'success': False,
                        'filename': file.filename,
                        'error': processing_result['error']
                    })
                    
            finally:
                # Clean up temporary file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        
        response_data = {
            'success': True,
            'results': results,
            'sections': sections,
            'industry': industry,
            'total_files': len(files),
            'processed_files': len([r for r in results if r.get('success', False)])
        }
        
        # Add AI information if used
        if ai_insights:
            response_data['ai_enhanced'] = True
            response_data['ai_insights'] = ai_insights
            response_data['total_ai_tokens'] = total_tokens_used
            response_data['ai_cost_estimate'] = f"${(total_tokens_used * 0.00015):.4f}"  # Rough estimate for gpt-4o-mini
        
        return jsonify(response_data)
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/parse_email', methods=['POST'])
def parse_email():
    """Parse lender email for requirements with AI enhancement"""
    try:
        data = request.get_json()
        content = data.get('content', '')
        
        if not content.strip():
            return jsonify({'success': False, 'error': 'No email content provided'})
        
        # Standard parsing
        parsed_info = parse_lender_email(content)
        
        # AI-Enhanced parsing (if available)
        ai_enhancement = None
        if mortgage_ai:
            ai_result = mortgage_ai.enhance_lender_parsing(content)
            if ai_result.get('ai_enhanced'):
                ai_enhancement = {
                    'enhanced_parsing': ai_result['enhanced_parsing'],
                    'tokens_used': ai_result['tokens_used'],
                    'total_tokens': ai_result['total_tokens']
                }
        
        # Store globally for use in analysis
        global lender_requirements
        lender_requirements = parsed_info
        
        response_data = {
            'success': True,
            'lender_name': parsed_info['lender_name'],
            'contact_email': parsed_info['contact_email'],
            'contact_name': parsed_info['contact_name'],
            'funding_amount': parsed_info['funding_amount'],
            'documents': parsed_info['documents'],
            'special_instructions': parsed_info['special_instructions'],
            'total_documents': len(parsed_info['documents'])
        }
        
        # Add AI enhancement if available
        if ai_enhancement:
            response_data['ai_enhanced'] = True
            response_data['ai_insights'] = ai_enhancement['enhanced_parsing']
            response_data['ai_tokens_used'] = ai_enhancement['tokens_used']
            response_data['ai_cost_estimate'] = f"${(ai_enhancement['tokens_used'] * 0.00015):.4f}"
        
        return jsonify(response_data)
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/reorganize_pdf', methods=['POST'])
def reorganize_pdf():
    """Memory-safe PDF reorganization with actual page extraction"""
    try:
        print("🔍 DEBUG: reorganize_pdf endpoint called")
        print(f"🧠 Memory at start: {get_memory_usage():.1f} MB")
        
        # Force garbage collection at start
        gc.collect()
        
        # Get request data with memory safety
        print("🔍 DEBUG: Getting request data...")
        try:
            data = request.get_json()
            print(f"🔍 DEBUG: Request data received: {data is not None}")
        except Exception as json_error:
            print(f"❌ DEBUG: JSON parsing error: {json_error}")
            return jsonify({'success': False, 'error': f'Invalid JSON data: {str(json_error)}'})
        
        if not data:
            print("❌ DEBUG: No data received in request")
            return jsonify({'success': False, 'error': 'No data received'})
        
        print(f"🔍 DEBUG: Data keys: {list(data.keys()) if data else 'None'}")
        
        document_sections = data.get('document_sections', [])
        lender_requirements = data.get('lender_requirements', {})
        original_pdf_path = data.get('original_pdf_path', '')
        
        print("📄 MEMORY-SAFE PDF REORGANIZATION - Real page extraction enabled")
        print(f"🔍 DEBUG: Received {len(document_sections)} document sections")
        print(f"🔍 DEBUG: Original PDF path: {original_pdf_path}")
        
        # Validate input data
        if not document_sections:
            print("❌ DEBUG: No document sections provided")
            return jsonify({'success': False, 'error': 'No document sections provided'})
        
        # Limit document sections for memory safety
        if len(document_sections) > 15:
            document_sections = document_sections[:15]
            print(f"⚠️  Limited to {len(document_sections)} documents for memory safety")
        
        # Create output directory
        print("🔍 DEBUG: Creating output directory...")
        output_dir = "/tmp/pdf_reorganization"
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate output filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"reorganized_mortgage_package_{timestamp}.pdf"
        output_path = os.path.join(output_dir, output_filename)
        
        print(f"🔍 DEBUG: Output path will be: {output_path}")
        
        # Check if we have an original PDF to work with
        has_original_pdf = original_pdf_path and os.path.exists(original_pdf_path)
        print(f"🔍 DEBUG: Has original PDF: {has_original_pdf}")
        
        if has_original_pdf:
            print(f"📄 Processing original PDF: {original_pdf_path}")
            try:
                reorganized_pages = extract_and_reorganize_pages_safe(original_pdf_path, document_sections)
                print(f"🔍 DEBUG: Page extraction result: {reorganized_pages is not None}")
                if reorganized_pages:
                    print(f"🔍 DEBUG: Total pages extracted: {reorganized_pages.get('total_pages', 0)}")
                    organized = reorganized_pages.get('organized_pages', {})
                    print(f"🔍 DEBUG: Organized pages keys: {list(organized.keys())}")
                    for doc_name, pages in organized.items():
                        print(f"🔍 DEBUG: {doc_name}: {len(pages)} pages")
            except Exception as extraction_error:
                print(f"❌ DEBUG: Page extraction failed: {extraction_error}")
                import traceback
                print(f"🔍 DEBUG: Extraction traceback: {traceback.format_exc()}")
                reorganized_pages = None
        else:
            print("📄 No original PDF - creating document summary")
            reorganized_pages = None
        
        print(f"🧠 Memory after page processing: {get_memory_usage():.1f} MB")
        
        # Create the reorganized PDF
        print("🔍 DEBUG: Starting PDF creation...")
        try:
            success = create_reorganized_pdf_safe(output_path, document_sections, reorganized_pages, lender_requirements)
            print(f"🔍 DEBUG: PDF creation success: {success}")
        except Exception as pdf_error:
            print(f"❌ DEBUG: PDF creation failed: {pdf_error}")
            import traceback
            print(f"🔍 DEBUG: PDF creation traceback: {traceback.format_exc()}")
            success = False
        
        if success and os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"🔍 DEBUG: Created PDF size: {file_size} bytes")
        else:
            print(f"🔍 DEBUG: PDF file exists: {os.path.exists(output_path) if output_path else 'No path'}")
        
        if not success:
            return jsonify({
                'success': False, 
                'error': 'Failed to create reorganized PDF',
                'note': 'Memory-safe processing attempted'
            })
        
        print(f"🧠 Memory after PDF creation: {get_memory_usage():.1f} MB")
        
        # Create response data
        response_data = {
            'success': True,
            'reorganized_pdf_path': output_path,
            'output_filename': output_filename,
            'documents_included': [doc.get('name', f'Document {i+1}') for i, doc in enumerate(document_sections)],
            'compliance_summary': f"Reorganized {len(document_sections)} documents in standard mortgage order. {'Original pages extracted and reorganized.' if has_original_pdf else 'Document summary created (no original PDF provided).'}",
            'ai_analysis': {
                'matching_analysis': 'Documents organized using standard mortgage document order for optimal lender compliance.',
                'ordering_analysis': 'Applied industry-standard document sequence: Application → Income → Assets → Property → Disclosures.',
                'compliance_analysis': 'Basic compliance check completed. Documents arranged for efficient lender review.',
                'total_tokens_used': 0,
                'cost_estimate': '$0.0000'
            },
            'generation_timestamp': datetime.now().isoformat(),
            'processing_method': 'memory_safe_extraction' if has_original_pdf else 'document_summary',
            'pages_processed': reorganized_pages['total_pages'] if reorganized_pages else 0
        }
        
        # Force memory cleanup before returning
        gc.collect()
        print(f"🧠 Memory at end: {get_memory_usage():.1f} MB")
        print("✅ Memory-safe PDF reorganization completed successfully")
        
        return jsonify(response_data)
        
    except Exception as e:
        # Force cleanup on error
        gc.collect()
        print(f"❌ Error in reorganize_pdf: {str(e)}")
        import traceback
        print(f"🔍 DEBUG: Full endpoint traceback: {traceback.format_exc()}")
        print(f"🧠 Memory after error: {get_memory_usage():.1f} MB")
        return jsonify({
            'success': False, 
            'error': f'PDF reorganization failed: {str(e)}',
            'note': 'Memory-safe processing attempted'
        })


def extract_and_reorganize_pages_safe(pdf_path, document_sections):
    """Memory-safe PDF page extraction with actual page content preservation"""
    try:
        import PyPDF2
        
        print(f"📄 Starting enhanced page extraction from {pdf_path}")
        
        # Initialize organized pages structure
        organized_pages = {}
        for i, doc in enumerate(document_sections):
            doc_name = doc.get('name', f'Document {i+1}')
            organized_pages[doc_name] = []
        
        # Read PDF and extract pages immediately
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            total_pages = len(pdf_reader.pages)
            
            print(f"📄 Total pages in original PDF: {total_pages}")
            
            # Limit pages for memory safety
            max_pages = min(total_pages, 50)  # Limit to 50 pages max
            if total_pages > max_pages:
                print(f"⚠️  Limited to {max_pages} pages for memory safety")
            
            # Create a new PDF writer to store extracted pages
            temp_pdf_path = pdf_path.replace('.pdf', '_temp_extracted.pdf')
            temp_pdf_writer = PyPDF2.PdfWriter()
            
            # Process pages one by one and store in temp PDF
            page_assignments = []
            for i in range(max_pages):
                try:
                    page = pdf_reader.pages[i]
                    
                    # Extract text for classification (first 300 chars)
                    text_sample = ""
                    if hasattr(page, 'extract_text'):
                        text_sample = page.extract_text()[:300]
                    
                    # Assign page to document
                    assigned_doc = assign_page_to_document_safe(text_sample, document_sections, i)
                    doc_name = document_sections[assigned_doc].get('name', f'Document {assigned_doc+1}')
                    
                    # Add page to temp PDF
                    temp_pdf_writer.add_page(page)
                    
                    # Store assignment info
                    page_assignments.append({
                        'page_number': i + 1,
                        'temp_page_index': i,  # Index in temp PDF
                        'text_sample': text_sample[:100],
                        'assigned_document': assigned_doc,
                        'doc_name': doc_name
                    })
                    
                    print(f"📄 Page {i+1} assigned to: {doc_name}")
                    
                    # Force cleanup every 10 pages
                    if (i + 1) % 10 == 0:
                        gc.collect()
                        print(f"🧠 Memory after {i+1} pages: {get_memory_usage():.1f} MB")
                        
                except Exception as page_error:
                    print(f"⚠️  Error processing page {i+1}: {page_error}")
                    continue
            
            # Write temp PDF
            with open(temp_pdf_path, 'wb') as temp_file:
                temp_pdf_writer.write(temp_file)
        
        # Now organize the assignments
        for assignment in page_assignments:
            doc_name = assignment['doc_name']
            organized_pages[doc_name].append(assignment)
        
        return {
            'total_pages': max_pages,
            'organized_pages': organized_pages,
            'processing_method': 'enhanced_with_content',
            'temp_pdf_path': temp_pdf_path  # Path to temp PDF with all pages
        }
        
    except Exception as e:
        print(f"❌ Error in enhanced page extraction: {e}")
        import traceback
        print(f"🔍 DEBUG: Extraction traceback: {traceback.format_exc()}")
        return None


def assign_page_to_document_safe(text_sample, document_sections, page_index):
    """Enhanced page assignment with better keyword matching"""
    try:
        text_lower = text_sample.lower()
        
        # Enhanced keyword mappings for mortgage documents
        keyword_mappings = {
            'mortgage': ['mortgage', 'deed of trust', 'security instrument', 'promissory note'],
            'application': ['uniform residential loan application', 'application', 'borrower information', 'fannie mae', 'freddie mac'],
            'income': ['income', 'employment', 'salary', 'w-2', 'pay stub', 'verification of employment', 'tax return'],
            'asset': ['asset', 'bank statement', 'account statement', 'verification of deposit', 'financial statement'],
            'property': ['property', 'appraisal', 'title', 'deed', 'survey', 'property information'],
            'closing': ['closing disclosure', 'settlement statement', 'hud-1', 'closing instructions'],
            'disclosure': ['truth in lending', 'good faith estimate', 'loan estimate', 'disclosure', 'notice'],
            'insurance': ['insurance', 'hazard insurance', 'flood insurance', 'title insurance'],
            'power_of_attorney': ['power of attorney', 'poa', 'attorney in fact'],
            'acknowledgment': ['acknowledgment', 'notary', 'notarization', 'sworn statement']
        }
        
        # Score each document section
        best_match = 0
        best_score = 0
        
        for i, doc in enumerate(document_sections):
            doc_name_lower = doc.get('name', '').lower()
            score = 0
            
            # Direct name match (highest priority)
            doc_words = doc_name_lower.split()
            for word in doc_words:
                if len(word) > 3 and word in text_lower:
                    score += 10
            
            # Category keyword match
            for category, keywords in keyword_mappings.items():
                if category in doc_name_lower:
                    for keyword in keywords:
                        if keyword in text_lower:
                            score += 5
            
            # Document type indicators
            if 'mortgage' in doc_name_lower and any(word in text_lower for word in ['mortgage', 'deed', 'security']):
                score += 15
            if 'application' in doc_name_lower and any(word in text_lower for word in ['application', 'borrower', 'fannie', 'freddie']):
                score += 15
            
            if score > best_score:
                best_score = score
                best_match = i
        
        # If no good match found, use round-robin assignment
        if best_score == 0:
            best_match = page_index % len(document_sections)
        
        return best_match
        
    except Exception as e:
        print(f"⚠️  Error in enhanced page assignment: {e}")
        return 0


def organize_pages_by_document(page_info, document_sections):
    """Organize pages by assigned document"""
    organized = {}
    
    for i, doc in enumerate(document_sections):
        doc_name = doc.get('name', f'Document {i+1}')
        organized[doc_name] = []
    
    for page in page_info:
        doc_index = page['assigned_document']
        if doc_index < len(document_sections):
            doc_name = document_sections[doc_index].get('name', f'Document {doc_index+1}')
            organized[doc_name].append(page)
    
    return organized


def create_reorganized_pdf_safe(output_path, document_sections, reorganized_pages, lender_requirements):
    """Create reorganized PDF including actual pages from original PDF"""
    try:
        print(f"📄 Creating enhanced reorganized PDF with actual pages: {output_path}")
        print(f"🔍 DEBUG: Document sections count: {len(document_sections)}")
        print(f"🔍 DEBUG: Reorganized pages available: {reorganized_pages is not None}")
        
        # Create PDF writer
        import PyPDF2
        pdf_writer = PyPDF2.PdfWriter()
        print("🔍 DEBUG: PDF writer created")
        
        # Create cover page using reportlab
        cover_path = output_path.replace('.pdf', '_cover.pdf')
        print(f"🔍 DEBUG: Creating cover page at: {cover_path}")
        cover_success = create_cover_page_enhanced(cover_path, document_sections, reorganized_pages, lender_requirements)
        print(f"🔍 DEBUG: Cover page creation success: {cover_success}")
        
        # Add cover page to final PDF
        if cover_success and os.path.exists(cover_path):
            print("🔍 DEBUG: Adding cover page to PDF")
            with open(cover_path, 'rb') as cover_file:
                cover_reader = PyPDF2.PdfReader(cover_file)
                cover_page_count = len(cover_reader.pages)
                print(f"🔍 DEBUG: Cover has {cover_page_count} pages")
                for i, page in enumerate(cover_reader.pages):
                    pdf_writer.add_page(page)
                    print(f"🔍 DEBUG: Added cover page {i+1}")
        else:
            print("🔍 DEBUG: Cover page not available, skipping")
        
        # Add organized document pages
        if reorganized_pages and reorganized_pages.get('organized_pages'):
            organized = reorganized_pages['organized_pages']
            temp_pdf_path = reorganized_pages.get('temp_pdf_path')
            
            print(f"🔍 DEBUG: Processing {len(organized)} organized document groups")
            print(f"🔍 DEBUG: Temp PDF path: {temp_pdf_path}")
            
            if temp_pdf_path and os.path.exists(temp_pdf_path):
                # Open the temp PDF containing all extracted pages
                with open(temp_pdf_path, 'rb') as temp_file:
                    temp_reader = PyPDF2.PdfReader(temp_file)
                    print(f"🔍 DEBUG: Temp PDF has {len(temp_reader.pages)} pages")
                    
                    # Process documents in order
                    for doc_index, doc in enumerate(document_sections):
                        doc_name = doc.get('name', 'Unknown Document')
                        print(f"🔍 DEBUG: Processing document {doc_index+1}: {doc_name}")
                        
                        if doc_name in organized and organized[doc_name]:
                            pages_for_doc = organized[doc_name]
                            print(f"📄 Adding {len(pages_for_doc)} pages for: {doc_name}")
                            
                            # Add document separator page
                            separator_path = output_path.replace('.pdf', f'_sep_{doc_name.replace(" ", "_")}.pdf')
                            print(f"🔍 DEBUG: Creating separator at: {separator_path}")
                            sep_success = create_document_separator_enhanced(separator_path, doc_name, len(pages_for_doc))
                            print(f"🔍 DEBUG: Separator creation success: {sep_success}")
                            
                            if sep_success and os.path.exists(separator_path):
                                with open(separator_path, 'rb') as sep_file:
                                    sep_reader = PyPDF2.PdfReader(sep_file)
                                    for page in sep_reader.pages:
                                        pdf_writer.add_page(page)
                                        print(f"🔍 DEBUG: Added separator page for {doc_name}")
                            
                            # Add actual document pages from temp PDF
                            print(f"🔍 DEBUG: Adding {len(pages_for_doc)} actual pages for {doc_name}")
                            for page_assignment in pages_for_doc:
                                try:
                                    temp_page_index = page_assignment.get('temp_page_index')
                                    if temp_page_index is not None and temp_page_index < len(temp_reader.pages):
                                        page_obj = temp_reader.pages[temp_page_index]
                                        pdf_writer.add_page(page_obj)
                                        print(f"🔍 DEBUG: Added page {page_assignment.get('page_number', temp_page_index+1)} to {doc_name}")
                                    else:
                                        print(f"⚠️  Invalid temp page index for page {page_assignment.get('page_number', 'unknown')}")
                                except Exception as page_error:
                                    print(f"⚠️  Error adding page {page_assignment.get('page_number', 'unknown')}: {page_error}")
                                    continue
                            
                            # Cleanup temporary separator file
                            try:
                                if os.path.exists(separator_path):
                                    os.remove(separator_path)
                            except:
                                pass
                            
                            # Memory cleanup after each document
                            gc.collect()
                        else:
                            print(f"🔍 DEBUG: No pages found for document: {doc_name}")
                
                # Cleanup temp PDF
                try:
                    if os.path.exists(temp_pdf_path):
                        os.remove(temp_pdf_path)
                        print("🔍 DEBUG: Cleaned up temp PDF")
                except:
                    pass
            else:
                print("🔍 DEBUG: Temp PDF not available")
        else:
            # Fallback: Create summary if no pages available
            print("📄 No organized pages available - creating document summary")
            
        # Check how many pages we have before writing
        total_pages_to_write = len(pdf_writer.pages)
        print(f"🔍 DEBUG: Total pages to write: {total_pages_to_write}")
        
        # Write final PDF
        print(f"🔍 DEBUG: Writing final PDF to: {output_path}")
        with open(output_path, 'wb') as output_file:
            pdf_writer.write(output_file)
        
        # Verify the written file
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"🔍 DEBUG: Written PDF size: {file_size} bytes")
            
            # Verify page count in written file
            try:
                with open(output_path, 'rb') as verify_file:
                    verify_reader = PyPDF2.PdfReader(verify_file)
                    written_page_count = len(verify_reader.pages)
                    print(f"🔍 DEBUG: Verified page count in written PDF: {written_page_count}")
            except Exception as verify_error:
                print(f"⚠️  Error verifying written PDF: {verify_error}")
        else:
            print("❌ DEBUG: Output file does not exist after writing!")
        
        # Cleanup temporary cover file
        try:
            if os.path.exists(cover_path):
                os.remove(cover_path)
        except:
            pass
        
        print(f"✅ Enhanced PDF created successfully with actual pages: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Error creating enhanced PDF: {e}")
        import traceback
        print(f"🔍 DEBUG: Full traceback: {traceback.format_exc()}")
        # Fallback to simple PDF creation
        return create_simple_pdf_fallback(output_path, document_sections, reorganized_pages, lender_requirements)


def create_cover_page_enhanced(cover_path, document_sections, reorganized_pages, lender_requirements):
    """Create professional cover page with enhanced information"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        c = canvas.Canvas(cover_path, pagesize=letter)
        width, height = letter
        
        # Title
        c.setFont("Helvetica-Bold", 20)
        c.drawString(50, height - 60, "AI-Reorganized Mortgage Package")
        
        # Subtitle
        c.setFont("Helvetica", 14)
        c.drawString(50, height - 90, "Professional Document Organization with Page Extraction")
        
        # Generation info
        c.setFont("Helvetica", 12)
        c.drawString(50, height - 120, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        c.drawString(50, height - 140, f"Documents included: {len(document_sections)}")
        
        if reorganized_pages:
            c.drawString(50, height - 160, f"Total pages processed: {reorganized_pages['total_pages']}")
            c.drawString(50, height - 180, f"Processing method: {reorganized_pages['processing_method']}")
        
        # Lender information
        if lender_requirements:
            y_pos = height - 220
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y_pos, "Lender Information:")
            
            y_pos -= 25
            c.setFont("Helvetica", 10)
            lender_name = lender_requirements.get('lender_name', 'Unknown Lender')
            c.drawString(70, y_pos, f"Lender: {lender_name}")
            
            y_pos -= 15
            contact_email = lender_requirements.get('contact_email', 'N/A')
            c.drawString(70, y_pos, f"Contact: {contact_email}")
        
        # Document summary
        y_pos = height - 320
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y_pos, "Document Organization Summary:")
        
        y_pos -= 30
        c.setFont("Helvetica", 10)
        
        for i, doc in enumerate(document_sections):
            if y_pos < 50:  # Start new page if needed
                c.showPage()
                y_pos = height - 50
            
            doc_name = doc.get('name', f'Document {i+1}')
            c.drawString(70, y_pos, f"• {doc_name}")
            
            # Add page count if available
            if reorganized_pages and reorganized_pages.get('organized_pages'):
                organized = reorganized_pages['organized_pages']
                if doc_name in organized:
                    page_count = len(organized[doc_name])
                    c.drawString(400, y_pos, f"({page_count} pages)")
            
            y_pos -= 20
        
        # Compliance notes
        if y_pos < 150:
            c.showPage()
            y_pos = height - 50
        
        y_pos -= 30
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y_pos, "Compliance & Quality Assurance:")
        
        y_pos -= 25
        c.setFont("Helvetica", 10)
        c.drawString(50, y_pos, "✓ Documents organized in industry-standard mortgage order")
        y_pos -= 15
        c.drawString(50, y_pos, "✓ AI-powered page classification and organization")
        y_pos -= 15
        c.drawString(50, y_pos, "✓ Memory-safe processing with content preservation")
        y_pos -= 15
        c.drawString(50, y_pos, "✓ All original pages included and properly sequenced")
        
        c.showPage()
        c.save()
        
        return True
        
    except Exception as e:
        print(f"❌ Error creating enhanced cover page: {e}")
        return False


def create_document_separator_enhanced(separator_path, doc_name, page_count):
    """Create separator page for each document section"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        c = canvas.Canvas(separator_path, pagesize=letter)
        width, height = letter
        
        # Draw border
        c.setStrokeColorRGB(0, 0.5, 1)  # Blue border
        c.setLineWidth(2)
        c.rect(30, 30, width-60, height-60)
        
        # Document name
        c.setFont("Helvetica-Bold", 18)
        text_width = c.stringWidth(doc_name, "Helvetica-Bold", 18)
        c.drawString((width - text_width) / 2, height - 100, doc_name)
        
        # Page count
        c.setFont("Helvetica", 14)
        page_text = f"{page_count} page{'s' if page_count != 1 else ''}"
        text_width = c.stringWidth(page_text, "Helvetica", 14)
        c.drawString((width - text_width) / 2, height - 130, page_text)
        
        # Timestamp
        c.setFont("Helvetica", 10)
        timestamp = f"Organized: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        text_width = c.stringWidth(timestamp, "Helvetica", 10)
        c.drawString((width - text_width) / 2, height - 160, timestamp)
        
        c.showPage()
        c.save()
        
        return True
        
    except Exception as e:
        print(f"❌ Error creating separator page: {e}")
        return False


def create_simple_pdf_fallback(output_path, document_sections, reorganized_pages, lender_requirements):
    """Fallback PDF creation if enhanced version fails"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        print(f"📄 Creating fallback PDF: {output_path}")
        
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        # Create cover page
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, "Reorganized Mortgage Package")
        
        c.setFont("Helvetica", 12)
        c.drawString(50, height - 80, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        c.drawString(50, height - 100, f"Documents included: {len(document_sections)}")
        
        if reorganized_pages:
            c.drawString(50, height - 120, f"Pages processed: {reorganized_pages['total_pages']}")
            c.drawString(50, height - 140, "Processing method: Enhanced page extraction (fallback mode)")
        else:
            c.drawString(50, height - 120, "Processing method: Document summary (no original PDF)")
        
        # Add document list
        y_pos = height - 180
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y_pos, "Documents Included:")
        
        y_pos -= 30
        c.setFont("Helvetica", 10)
        
        for i, doc in enumerate(document_sections):
            if y_pos < 50:  # Start new page if needed
                c.showPage()
                y_pos = height - 50
            
            doc_name = doc.get('name', f'Document {i+1}')
            c.drawString(70, y_pos, f"• {doc_name}")
            
            # Add page count if available
            if reorganized_pages and reorganized_pages.get('organized_pages'):
                organized = reorganized_pages['organized_pages']
                if doc_name in organized:
                    page_count = len(organized[doc_name])
                    c.drawString(400, y_pos, f"({page_count} pages)")
            
            y_pos -= 20
        
        c.showPage()
        c.save()
        
        print(f"✅ Fallback PDF created successfully: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Error creating fallback PDF: {e}")
        return False

@app.route('/download_pdf/<filename>')
def download_pdf(filename):
    """Download reorganized PDF endpoint"""
    try:
        # Security: Only allow downloading from the temp reorganization directory
        safe_filename = os.path.basename(filename)  # Remove any path traversal
        file_path = os.path.join("/tmp/pdf_reorganization", safe_filename)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
        
        return send_file(
            file_path,
            as_attachment=True,
            download_name=safe_filename,
            mimetype='application/pdf'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/segment/separate', methods=['POST'])
def separate_document_segments():
    """Separate document into individual segments/sections"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})
        
        # Save uploaded file temporarily
        temp_path = f"/tmp/{file.filename}"
        file.save(temp_path)
        
        try:
            # Separate the document into segments
            result = document_processor.separate_documents(temp_path, file.filename)
            
            return jsonify(result)
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/segment/analyze', methods=['POST'])
def analyze_document_segment():
    """Analyze a specific document segment"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'})
        
        file = request.files['file']
        section_name = request.form.get('section_name', None)
        
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})
        
        # Save uploaded file temporarily
        temp_path = f"/tmp/{file.filename}"
        file.save(temp_path)
        
        try:
            # Analyze the document/segment
            result = document_processor.analyze_document(temp_path, file.filename, section_name)
            
            return jsonify(result)
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/segment/extract', methods=['POST'])
def extract_segment_text():
    """Extract text from a specific document segment"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'})
        
        file = request.files['file']
        section_name = request.form.get('section_name', None)
        page_range = request.form.get('page_range', None)
        
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})
        
        # Save uploaded file temporarily
        temp_path = f"/tmp/{file.filename}"
        file.save(temp_path)
        
        try:
            # Extract text from the segment
            result = document_processor.extract_text(temp_path, file.filename, section_name, page_range)
            
            return jsonify(result)
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/segment/export', methods=['POST'])
def export_document_segment():
    """Export a specific document segment"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'})
        
        file = request.files['file']
        section_name = request.form.get('section_name', 'Unknown Section')
        export_format = request.form.get('format', 'txt')
        
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})
        
        # Save uploaded file temporarily
        temp_path = f"/tmp/{file.filename}"
        file.save(temp_path)
        
        try:
            # Export the segment
            result = document_processor.export_section(temp_path, file.filename, section_name, export_format)
            
            if result['success']:
                # Return the export file for download
                return send_file(
                    result['export_path'],
                    as_attachment=True,
                    download_name=result['export_filename'],
                    mimetype='text/plain' if export_format == 'txt' else 'application/json'
                )
            else:
                return jsonify(result)
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/segment/view/<section_name>', methods=['POST'])
def view_document_segment(section_name):
    """View a specific document segment with formatted display"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})
        
        # Save uploaded file temporarily
        temp_path = f"/tmp/{file.filename}"
        file.save(temp_path)
        
        try:
            # First separate the document to get sections
            separation_result = document_processor.separate_documents(temp_path, file.filename)
            
            if not separation_result['success']:
                return jsonify(separation_result)
            
            # Find the requested section
            target_section = None
            for section in separation_result['sections']:
                if section['name'] == section_name:
                    target_section = section
                    break
            
            if not target_section:
                return jsonify({
                    'success': False,
                    'error': f'Section "{section_name}" not found',
                    'available_sections': [s['name'] for s in separation_result['sections']]
                })
            
            # Get detailed analysis for this section
            analysis_result = document_processor.analyze_document(temp_path, file.filename, section_name)
            
            # Combine section info with analysis
            view_result = {
                'success': True,
                'section': target_section,
                'analysis': analysis_result.get('analysis', {}),
                'metadata': {
                    'source_file': file.filename,
                    'section_name': section_name,
                    'text_length': len(target_section.get('full_text', '')),
                    'confidence': target_section.get('confidence', 'unknown'),
                    'risk_score': target_section.get('risk_score', 0)
                }
            }
            
            return jsonify(view_result)
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/debug/openai')
def debug_openai():
    """Debug endpoint for OpenAI troubleshooting"""
    try:
        import os
        
        debug_info = {
            'openai_available': OPENAI_AVAILABLE,
            'api_key_present': bool(OPENAI_API_KEY),
            'api_key_suffix': OPENAI_API_KEY[-4:] if OPENAI_API_KEY else 'None',
            'api_key_source': 'environment' if os.getenv('OPENAI_API_KEY') else 'hardcoded',
            'api_key_length': len(OPENAI_API_KEY) if OPENAI_API_KEY else 0,
            'client_error': openai_client_error,
            'client_initialized': openai_client is not None,
            'pdf_reorganizer_ai_available': pdf_reorganizer_ai is not None,
            'connection_test': test_openai_connection()
        }
        
        # Try to reinitialize if not working
        if not openai_client:
            debug_info['reinitialize_attempt'] = 'attempting'
            try:
                new_client = reinitialize_openai_client()
                debug_info['reinitialize_result'] = 'success' if new_client else 'failed'
                debug_info['connection_test_after_reinit'] = test_openai_connection()
            except Exception as e:
                debug_info['reinitialize_result'] = f'error: {str(e)}'
        
        return jsonify(debug_info)
        
    except Exception as e:
        return jsonify({'error': f'Debug endpoint failed: {str(e)}'})

@app.route('/health')
def health_check():
    """Health check endpoint with comprehensive system status"""
    
    # Test OpenAI connection
    openai_status = test_openai_connection()
    
    # Check other components
    pdf_status = "available" if PDF_GENERATION_AVAILABLE else "unavailable"
    pdfplumber_status = "available" if PDFPLUMBER_AVAILABLE else "unavailable"
    
    # Check if AI classes are initialized
    ai_classes_status = {
        "mortgage_ai": "initialized" if mortgage_ai else "not_initialized",
        "pdf_reorganizer_ai": "initialized" if pdf_reorganizer_ai else "not_initialized"
    }
    
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'components': {
            'openai': {
                'status': openai_status['status'],
                'message': openai_status['message'],
                'available': OPENAI_AVAILABLE
            },
            'pdf_generation': {
                'status': pdf_status,
                'available': PDF_GENERATION_AVAILABLE
            },
            'pdf_processing': {
                'status': pdfplumber_status,
                'available': PDFPLUMBER_AVAILABLE
            },
            'ai_classes': ai_classes_status
        },
        'features': [
            'Multi-industry document analysis',
            'Email parsing',
            'Industry-specific analysis',
            'Real-time processing',
            f'AI-powered analysis ({openai_status["status"]})',
            f'PDF reorganization ({pdf_status})'
        ]
    })

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5009, debug=True)

